import streamlit as st
import math
import datetime
import pandas as pd
import base64
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
import os
import traceback

st.set_page_config(page_title="Professional Engineering Tools", page_icon="🔌", layout="wide")

# ========== PAKISTAN TIME HELPER FUNCTIONS ==========
def get_pakistan_time():
    """Get current Pakistan time (UTC+5)"""
    pakistan_time = datetime.utcnow() + timedelta(hours=5)
    return pakistan_time

def format_pakistan_datetime():
    """Format Pakistan time for reports"""
    return get_pakistan_time().strftime("%Y-%m-%d %H:%M:%S")

def format_pakistan_date():
    """Format Pakistan date only for reports"""
    return get_pakistan_time().strftime("%Y-%m-%d")

# ========== HELPER FUNCTIONS FOR FORMATTING ==========
def format_cable_arrangement(arrangement):
    formats = {
        'bunched_in_air_surface_enclosed': 'Bunched in air / surface / enclosed',
        'single_layer_wall_floor': 'Single layer on wall or floor',
        'single_layer_perforated_tray': 'Single layer on perforated tray',
        'single_layer_ladder_cleats': 'Single layer on ladder / cleats',
        'direct_buried': 'Direct Buried in Ground',
        'buried_ducts': 'Buried in Ducts'
    }
    return formats.get(arrangement, arrangement.replace('_', ' ').title())

def format_cable_clearance(clearance):
    formats = {
        'touching': 'Touching (0 clearance)',
        'one_diameter': 'One cable diameter spacing',
        'clearance_0_125m': '0.125 m clearance',
        'clearance_0_25m': '0.25 m clearance',
        'clearance_0_5m': '0.5 m clearance',
        'clearance_1_0m': '1.0 m clearance'
    }
    return formats.get(clearance, clearance.replace('_', ' ').title())

def format_cable_formation(formation):
    formats = {'flat': 'Flat', 'trefoil': 'Trefoil', 'spaced': 'Spaced'}
    return formats.get(formation, formation.title())

def format_cable_type(cable_type):
    formats = {
        'single_core_non_armoured': 'Single core non-armoured',
        'multi_core_non_armoured': 'Multi core non-armoured',
        'single_core_armoured': 'Single core armoured',
        'multi_core_armoured': 'Multi core armoured'
    }
    return formats.get(cable_type, cable_type.replace('_', ' ').title())

def format_insulation_type(insulation_type):
    if insulation_type == 'PVC_70':
        return 'PVC 70°C'
    elif insulation_type == 'XLPE_90':
        return 'XLPE 90°C'
    return insulation_type.replace('_', ' ')

def format_load_type(load_type):
    return load_type.capitalize()

def format_installation_method(method):
    formats = {
        'B': 'Method B - Enclosed in conduit on wall',
        'C': 'Method C - Clipped direct / on tray / in free air',
        'D': 'Method D - Buried in ducts',
        'D_direct': 'Method D - Direct buried in ground',
        'E': 'Method E - On open perforated cable tray',
        'F': 'Method F - In free air (trefoil/flat)',
        'G': 'Method G - In free air (spaced)'
    }
    return formats.get(method, method)

def get_table_config_description(config_key, cable_type):
    descriptions = {
        'B2': '2 cables, single-phase a.c. or d.c.',
        'B34': '3 or 4 cables, three-phase a.c.',
        'C2': '2 cables, single-phase a.c. or d.c.',
        'C34': '3 or 4 cables, three-phase a.c.',
        'F2': '2 cables, flat and touching, single-phase a.c. or d.c.',
        'F34_flat': '3 or 4 cables, flat and touching, three-phase a.c.',
        'F34_trefoil': '3 cables, trefoil formation, three-phase a.c.',
        'G2': '2 cables, spaced by one diameter, single-phase a.c. or d.c.',
        'G34': '3 cables, spaced by one diameter, three-phase a.c.',
        'E2': '1 two-core cable, single-phase a.c. or d.c.',
        'E34': '1 three or four-core cable, three-phase a.c.',
        'F2_flat': '2 cables, flat and touching, single-phase a.c. or d.c.',
        'D2': '1 two-core cable, single-phase a.c. or d.c. (buried)',
        'D34': '1 three or four-core cable, three-phase a.c. (buried)'
    }
    return descriptions.get(config_key, config_key)

st.markdown("""
<style>
    :root { --primary: #1E3A8A; --secondary: #00A86B; --light-bg: #F0F8FF; --card-bg: #FFFFFF; --text-dark: #1a1a1a; --text-muted: #4a4a4a; }
    .report-header { background: linear-gradient(135deg, var(--primary) 0%, #3B5BA6 100%); color: white; padding: 25px; border-radius: 12px; text-align: center; margin-bottom: 25px; font-size: 32px !important; font-weight: bold; }
    .info-box { background: linear-gradient(135deg, #E8F4FD 0%, #D6ECFA 100%); color: var(--text-dark) !important; padding: 15px; border-radius: 8px; border-left: 5px solid var(--primary); margin: 10px 0; font-size: 18px !important; }
    .info-box * { color: var(--text-dark) !important; }
    .info-box h4 { color: var(--primary) !important; }
    .calc-step { background-color: #F8F9FA !important; padding: 15px; border-radius: 8px; border-left: 5px solid #00A86B; margin: 10px 0; box-shadow: 0 2px 4px rgba(0,0,0,0.1); font-size: 18px !important; color: var(--text-dark) !important; }
    .calc-step * { color: var(--text-dark) !important; }
    .calc-step h4 { color: var(--primary) !important; margin-top: 0; margin-bottom: 10px; }
    .calc-step p { color: var(--text-dark) !important; margin: 5px 0; }
    .calc-step b { color: var(--primary) !important; }
    .result-card { background: linear-gradient(135deg, #1E3A8A 0%, #3B5BA6 100%); color: white !important; padding: 25px; border-radius: 12px; margin: 20px 0; }
    .download-btn { display: inline-block; padding: 14px 28px; margin: 10px; color: white !important; text-decoration: none; border-radius: 8px; font-size: 18px !important; font-weight: bold; text-align: center; background: linear-gradient(135deg, #1E3A8A 0%, #3B5BA6 100%); }
    .load-type-badge { display: inline-block; padding: 5px 12px; border-radius: 12px; font-size: 14px !important; font-weight: bold; }
    .continuous-badge { background-color: #00A86B; color: white !important; }
    .intermittent-badge { background-color: #FFC107; color: #1a1a1a !important; }
    .standby-badge { background-color: #DC3545; color: white !important; }
    .stTabs [data-baseweb="tab"] { font-size: 18px !important; font-weight: 600 !important; }
    .stDataFrame { color: var(--text-dark) !important; }
    .stDataFrame table { color: var(--text-dark) !important; }
    .stDataFrame th { background: linear-gradient(135deg, var(--primary) 0%, #3B5BA6 100%) !important; color: white !important; }
    .stDataFrame td { color: var(--text-dark) !important; background-color: white !important; }
    div[data-testid="stMetricValue"] { color: var(--primary) !important; font-size: 28px !important; }
    div[data-testid="stMetricLabel"] { color: var(--text-muted) !important; }
    .largest-equipment { background: linear-gradient(135deg, #E8F5E9 0%, #C8E6C9 100%); padding: 20px; border-radius: 10px; border-left: 6px solid #00A86B; margin: 15px 0; color: var(--text-dark) !important; }
    .largest-equipment * { color: var(--text-dark) !important; }
    .largest-equipment h3 { color: #006B3C !important; }
    .largest-equipment .value { font-size: 20px !important; font-weight: bold; color: #006B3C !important; }
    .formula-box { background: linear-gradient(135deg, #F8F9FA 0%, #E9ECEF 100%); padding: 20px; border-radius: 10px; border-left: 6px solid var(--secondary); margin: 15px 0; color: var(--text-dark) !important; }
    .formula-box * { color: var(--text-dark) !important; }
    .formula-box h4 { color: var(--primary) !important; }
</style>
""", unsafe_allow_html=True)

CB_RATINGS = [6, 10, 16, 20, 25, 32, 40, 50, 63, 80, 100, 125, 160, 200, 250, 315, 400, 500, 630, 800, 1000, 1250, 1600]

BREAKER_TYPES = {
    'MCB': {'min': 6, 'max': 125, 'standard': 'IEC 60898', 'application': 'Miniature Circuit Breaker - For final circuits'},
    'MCCB': {'min': 125, 'max': 1600, 'standard': 'IEC 60947-2', 'application': 'Moulded Case Circuit Breaker - For distribution'},
    'ACB': {'min': 1600, 'max': 6300, 'standard': 'IEC 60947-2', 'application': 'Air Circuit Breaker - For main incomers'}
}

MANUFACTURERS = {
    'Schneider Electric': {'MCB': 'Acti9 series', 'MCCB': 'EasyPact EVC series', 'ACB': 'MasterPact MTZ series'},
    'Siemens': {'MCB': '5SY series', 'MCCB': '3VA series', 'ACB': '3WL series'},
    'ABB': {'MCB': 'S200 series', 'MCCB': 'Tmax XT series', 'ACB': 'Emax 2 series'},
}

LOAD_TYPE_FACTORS = {
    'Continuous': {'diversity': 1.0, 'description': 'Continuous (100%) - Full time operation', 'cb_factor': 1.25, 'color': '#00A86B'},
    'Intermittent': {'diversity': 0.3, 'description': 'Intermittent (30%) - Cyclic operation', 'cb_factor': 1.25, 'color': '#FFC107'},
    'Standby': {'diversity': 0.1, 'description': 'Stand-by (10%) - Emergency/backup only', 'cb_factor': 1.25, 'color': '#DC3545'}
}

# ========== TEMPERATURE DERATING FACTORS ==========
TEMPERATURE_FACTORS_AIR = {70: {25: 1.03, 30: 1.00, 35: 0.94, 40: 0.87, 45: 0.79, 50: 0.71, 55: 0.61}, 90: {25: 1.02, 30: 1.00, 35: 0.96, 40: 0.91, 45: 0.87, 50: 0.82, 55: 0.76}}
TEMPERATURE_FACTORS_GROUND = {70: {10: 1.10, 15: 1.05, 20: 1.00, 25: 0.95, 30: 0.89, 35: 0.84, 40: 0.77, 45: 0.71}, 90: {10: 1.07, 15: 1.04, 20: 1.00, 25: 0.96, 30: 0.93, 35: 0.89, 40: 0.85, 45: 0.80}}

def get_temperature_factor(insulation_temp, ambient_temp, installation='air'):
    try:
        if installation in ['buried', 'duct', 'trench', 'ground', 'D', 'D_direct']:
            factors = TEMPERATURE_FACTORS_GROUND.get(insulation_temp, TEMPERATURE_FACTORS_GROUND[90])
        else:
            factors = TEMPERATURE_FACTORS_AIR.get(insulation_temp, TEMPERATURE_FACTORS_AIR[90])
        temps = sorted(factors.keys())
        closest_temp = min(temps, key=lambda x: abs(x - ambient_temp))
        return factors[closest_temp]
    except:
        return 1.0

# ========== GROUPING FACTORS ==========
GROUPING_FACTORS_AIR = {
    'bunched_in_air_surface_enclosed': {1: 1.00, 2: 0.80, 3: 0.70, 4: 0.65, 5: 0.60, 6: 0.57, 7: 0.54, 8: 0.52, 9: 0.50, 12: 0.45, 16: 0.41, 20: 0.38},
    'single_layer_wall_floor': {1: 1.00, 2: 0.85, 3: 0.79, 4: 0.75, 5: 0.73, 6: 0.72, 7: 0.72, 8: 0.71, 9: 0.70, 12: 0.70, 16: 0.70, 20: 0.70},
    'single_layer_perforated_tray': {1: 1.00, 2: 0.88, 3: 0.82, 4: 0.77, 5: 0.75, 6: 0.73, 7: 0.73, 8: 0.72, 9: 0.72, 12: 0.72, 16: 0.72, 20: 0.72},
    'single_layer_ladder_cleats': {1: 1.00, 2: 0.87, 3: 0.82, 4: 0.80, 5: 0.80, 6: 0.79, 7: 0.79, 8: 0.78, 9: 0.78, 12: 0.78, 16: 0.78, 20: 0.78},
}

GROUPING_FACTORS_BURIED_DIRECT = {
    'touching': {1: 1.00, 2: 0.75, 3: 0.65, 4: 0.60, 5: 0.55, 6: 0.50},
    'one_diameter': {1: 1.00, 2: 0.80, 3: 0.70, 4: 0.60, 5: 0.55, 6: 0.55},
    'clearance_0_125m': {1: 1.00, 2: 0.85, 3: 0.75, 4: 0.70, 5: 0.65, 6: 0.60},
    'clearance_0_25m': {1: 1.00, 2: 0.90, 3: 0.80, 4: 0.75, 5: 0.70, 6: 0.70},
    'clearance_0_5m': {1: 1.00, 2: 0.90, 3: 0.85, 4: 0.80, 5: 0.80, 6: 0.80},
}

GROUPING_FACTORS_BURIED_DUCTS_MULTI = {
    'touching': {1: 1.00, 2: 0.85, 3: 0.75, 4: 0.70, 5: 0.65, 6: 0.60},
    'clearance_0_25m': {1: 1.00, 2: 0.90, 3: 0.85, 4: 0.80, 5: 0.80, 6: 0.80},
    'clearance_0_5m': {1: 1.00, 2: 0.95, 3: 0.90, 4: 0.85, 5: 0.85, 6: 0.85},
    'clearance_1_0m': {1: 1.00, 2: 0.95, 3: 0.95, 4: 0.90, 5: 0.90, 6: 0.90},
}

GROUPING_FACTORS_BURIED_DUCTS_SINGLE = {
    'touching': {1: 1.00, 2: 0.80, 3: 0.70, 4: 0.65, 5: 0.60, 6: 0.60},
    'clearance_0_25m': {1: 1.00, 2: 0.90, 3: 0.80, 4: 0.75, 5: 0.70, 6: 0.70},
    'clearance_0_5m': {1: 1.00, 2: 0.90, 3: 0.85, 4: 0.80, 5: 0.80, 6: 0.80},
    'clearance_1_0m': {1: 1.00, 2: 0.95, 3: 0.90, 4: 0.90, 5: 0.90, 6: 0.95},
}

def get_grouping_factor(num_cables, arrangement, installation_type='air', clearance='touching', is_single_core=True):
    if installation_type in ['buried', 'duct', 'ground', 'D', 'D_direct']:
        if arrangement == 'direct_buried':
            if clearance in GROUPING_FACTORS_BURIED_DIRECT:
                factors = GROUPING_FACTORS_BURIED_DIRECT[clearance]
            else:
                factors = GROUPING_FACTORS_BURIED_DIRECT['touching']
        elif arrangement == 'buried_ducts':
            if is_single_core:
                if clearance in GROUPING_FACTORS_BURIED_DUCTS_SINGLE:
                    factors = GROUPING_FACTORS_BURIED_DUCTS_SINGLE[clearance]
                else:
                    factors = GROUPING_FACTORS_BURIED_DUCTS_SINGLE['touching']
            else:
                if clearance in GROUPING_FACTORS_BURIED_DUCTS_MULTI:
                    factors = GROUPING_FACTORS_BURIED_DUCTS_MULTI[clearance]
                else:
                    factors = GROUPING_FACTORS_BURIED_DUCTS_MULTI['touching']
        else:
            factors = GROUPING_FACTORS_BURIED_DIRECT['touching']
    else:
        if arrangement in GROUPING_FACTORS_AIR:
            factors = GROUPING_FACTORS_AIR[arrangement]
        else:
            factors = GROUPING_FACTORS_AIR['bunched_in_air_surface_enclosed']
    
    available = sorted(factors.keys())
    if num_cables in factors:
        return factors[num_cables]
    elif num_cables > max(available):
        return factors[max(available)]
    else:
        closest = min(available, key=lambda x: abs(x - num_cables))
        return factors[closest]

# ========== DEPTH FACTORS ==========
def get_depth_factor(depth_m, installation='D_direct', is_single_core=True):
    if depth_m == 0.8:
        return 1.00
    if depth_m < 0.5:
        depth_m = 0.5
    if depth_m > 3.0:
        depth_m = 3.0
    
    if installation == 'D':
        if is_single_core:
            factors = {0.5: 1.04, 0.6: 1.02, 0.8: 1.00, 1.0: 0.98, 1.25: 0.96, 1.5: 0.95, 1.75: 0.94, 2.0: 0.93, 2.5: 0.91, 3.0: 0.90}
        else:
            factors = {0.5: 1.03, 0.6: 1.02, 0.8: 1.00, 1.0: 0.99, 1.25: 0.97, 1.5: 0.96, 1.75: 0.95, 2.0: 0.94, 2.5: 0.93, 3.0: 0.92}
    else:
        if is_single_core:
            factors = {0.5: 1.04, 0.6: 1.02, 0.8: 1.00, 1.0: 0.98, 1.25: 0.96, 1.5: 0.95, 1.75: 0.94, 2.0: 0.93, 2.5: 0.91, 3.0: 0.90}
        else:
            factors = {0.5: 1.04, 0.6: 1.03, 0.8: 1.00, 1.0: 0.98, 1.25: 0.96, 1.5: 0.95, 1.75: 0.94, 2.0: 0.93, 2.5: 0.91, 3.0: 0.90}
    
    depths = sorted(factors.keys())
    closest_depth = min(depths, key=lambda x: abs(x - depth_m))
    return factors[closest_depth]

# ========== SOIL RESISTIVITY FACTORS ==========
def get_soil_resistivity_factor(soil_resistivity, installation='D_direct', is_single_core=True):
    if soil_resistivity == 1.5:
        return 1.00
    if soil_resistivity < 0.7:
        soil_resistivity = 0.7
    if soil_resistivity > 3.0:
        soil_resistivity = 3.0
    
    if installation == 'D':
        if is_single_core:
            factors = {0.7: 1.22, 0.8: 1.19, 0.9: 1.15, 1.0: 1.12, 1.5: 1.00, 2.0: 0.91, 2.5: 0.84, 3.0: 0.78}
        else:
            factors = {0.7: 1.15, 0.8: 1.13, 0.9: 1.11, 1.0: 1.09, 1.5: 1.00, 2.0: 0.94, 2.5: 0.88, 3.0: 0.83}
    else:
        if is_single_core:
            factors = {0.7: 1.33, 0.8: 1.27, 0.9: 1.22, 1.0: 1.17, 1.5: 1.00, 2.0: 0.89, 2.5: 0.81, 3.0: 0.74}
        else:
            factors = {0.7: 1.26, 0.8: 1.21, 0.9: 1.18, 1.0: 1.14, 1.5: 1.00, 2.0: 0.90, 2.5: 0.83, 3.0: 0.77}
    
    resistivities = sorted(factors.keys())
    closest_res = min(resistivities, key=lambda x: abs(x - soil_resistivity))
    return factors[closest_res]

def calculate_short_circuit_current(size_mm2, insulation_type, duration_s=1.0, conductor_material='Copper'):
    if conductor_material == 'Copper':
        K = 226
        β = 234.5
    else:
        K = 148
        β = 228
    if insulation_type == 'PVC':
        θi = 70
        θf = 160 if size_mm2 <= 300 else 140
    else:
        θi = 90
        θf = 250
    first_term = K * size_mm2 / math.sqrt(duration_s)
    log_term = math.log((θf + β) / (θi + β))
    return first_term * math.sqrt(log_term), K, θi, θf

# ========== ACCURATE VOLTAGE DROP DATABASE BASED ON BS 7671 TABLES ==========

# MULTI CORE NON-ARMOURED (Table 4E2B)
XLPE_90_MULTI_NON_ARMOURED = {
    1.0: {'voltage_drop': {'dc_mv': 46, 'single_phase_mv': 46, 'three_phase_mv': 40}},
    1.5: {'voltage_drop': {'dc_mv': 31, 'single_phase_mv': 31, 'three_phase_mv': 27}},
    2.5: {'voltage_drop': {'dc_mv': 19, 'single_phase_mv': 19, 'three_phase_mv': 16}},
    4.0: {'voltage_drop': {'dc_mv': 12, 'single_phase_mv': 12, 'three_phase_mv': 10}},
    6.0: {'voltage_drop': {'dc_mv': 7.9, 'single_phase_mv': 7.9, 'three_phase_mv': 6.8}},
    10.0: {'voltage_drop': {'dc_mv': 4.7, 'single_phase_mv': 4.7, 'three_phase_mv': 4.0}},
    16.0: {'voltage_drop': {'dc_mv': 2.9, 'single_phase_mv': 2.9, 'three_phase_mv': 2.5}},
    25.0: {'voltage_drop': {'single_phase_R': 1.85, 'single_phase_X': 0.16, 'single_phase_Z': 1.90,
                           'three_phase_R': 1.6, 'three_phase_X': 0.14, 'three_phase_Z': 1.65}},
    35.0: {'voltage_drop': {'single_phase_R': 1.35, 'single_phase_X': 0.155, 'single_phase_Z': 1.35,
                           'three_phase_R': 1.15, 'three_phase_X': 0.135, 'three_phase_Z': 1.15}},
    50.0: {'voltage_drop': {'single_phase_R': 0.99, 'single_phase_X': 0.155, 'single_phase_Z': 1.00,
                           'three_phase_R': 0.86, 'three_phase_X': 0.135, 'three_phase_Z': 0.87}},
    70.0: {'voltage_drop': {'single_phase_R': 0.67, 'single_phase_X': 0.15, 'single_phase_Z': 0.69,
                           'three_phase_R': 0.59, 'three_phase_X': 0.13, 'three_phase_Z': 0.60}},
    95.0: {'voltage_drop': {'single_phase_R': 0.50, 'single_phase_X': 0.15, 'single_phase_Z': 0.52,
                           'three_phase_R': 0.43, 'three_phase_X': 0.13, 'three_phase_Z': 0.45}},
    120.0: {'voltage_drop': {'single_phase_R': 0.40, 'single_phase_X': 0.145, 'single_phase_Z': 0.42,
                            'three_phase_R': 0.34, 'three_phase_X': 0.13, 'three_phase_Z': 0.37}},
    150.0: {'voltage_drop': {'single_phase_R': 0.32, 'single_phase_X': 0.145, 'single_phase_Z': 0.35,
                            'three_phase_R': 0.28, 'three_phase_X': 0.125, 'three_phase_Z': 0.30}},
    185.0: {'voltage_drop': {'single_phase_R': 0.26, 'single_phase_X': 0.145, 'single_phase_Z': 0.29,
                            'three_phase_R': 0.22, 'three_phase_X': 0.125, 'three_phase_Z': 0.26}},
    240.0: {'voltage_drop': {'single_phase_R': 0.20, 'single_phase_X': 0.14, 'single_phase_Z': 0.24,
                            'three_phase_R': 0.175, 'three_phase_X': 0.125, 'three_phase_Z': 0.21}},
    300.0: {'voltage_drop': {'single_phase_R': 0.16, 'single_phase_X': 0.14, 'single_phase_Z': 0.21,
                            'three_phase_R': 0.14, 'three_phase_X': 0.12, 'three_phase_Z': 0.185}},
    400.0: {'voltage_drop': {'single_phase_R': 0.13, 'single_phase_X': 0.14, 'single_phase_Z': 0.19,
                            'three_phase_R': 0.115, 'three_phase_X': 0.12, 'three_phase_Z': 0.165}},
}

# SINGLE CORE NON-ARMOURED (Table 4E1B)
XLPE_90_SINGLE_NON_ARMOURED = {
    1.5: {'voltage_drop': {'single_phase_mv': 31, 'three_phase_mv': 27}},
    2.5: {'voltage_drop': {'single_phase_mv': 19, 'three_phase_mv': 16}},
    4.0: {'voltage_drop': {'single_phase_mv': 12, 'three_phase_mv': 10}},
    6.0: {'voltage_drop': {'single_phase_mv': 7.9, 'three_phase_mv': 6.8}},
    10.0: {'voltage_drop': {'single_phase_mv': 4.7, 'three_phase_mv': 4.0}},
    16.0: {'voltage_drop': {'single_phase_mv': 2.9, 'three_phase_mv': 2.5}},
    25.0: {'voltage_drop': {'single_phase_R': 1.85, 'single_phase_X': 0.31, 'three_phase_R': 1.85, 'three_phase_X': 0.27}},
    35.0: {'voltage_drop': {'single_phase_R': 1.35, 'single_phase_X': 0.29, 'three_phase_R': 1.35, 'three_phase_X': 0.25}},
    50.0: {'voltage_drop': {'single_phase_R': 1.00, 'single_phase_X': 0.29, 'three_phase_R': 0.87, 'three_phase_X': 0.25}},
    70.0: {'voltage_drop': {'single_phase_R': 0.70, 'single_phase_X': 0.28, 'three_phase_R': 0.60, 'three_phase_X': 0.24}},
    95.0: {'voltage_drop': {'single_phase_R': 0.51, 'single_phase_X': 0.27, 'three_phase_R': 0.44, 'three_phase_X': 0.23}},
    120.0: {'voltage_drop': {'single_phase_R': 0.41, 'single_phase_X': 0.26, 'three_phase_R': 0.35, 'three_phase_X': 0.23}},
    150.0: {'voltage_drop': {'single_phase_R': 0.33, 'single_phase_X': 0.26, 'three_phase_R': 0.29, 'three_phase_X': 0.23}},
    185.0: {'voltage_drop': {'single_phase_R': 0.27, 'single_phase_X': 0.26, 'three_phase_R': 0.23, 'three_phase_X': 0.23}},
    240.0: {'voltage_drop': {'single_phase_R': 0.21, 'single_phase_X': 0.26, 'three_phase_R': 0.185, 'three_phase_X': 0.22}},
    300.0: {'voltage_drop': {'single_phase_R': 0.175, 'single_phase_X': 0.25, 'three_phase_R': 0.15, 'three_phase_X': 0.22}},
    400.0: {'voltage_drop': {'single_phase_R': 0.140, 'single_phase_X': 0.25, 'three_phase_R': 0.125, 'three_phase_X': 0.22}},
}

# SINGLE CORE ARMOURED (Table 4E3B)
XLPE_90_SINGLE_ARMOURED = {
    50.0: {'voltage_drop': {'single_phase_touching_R': 0.99, 'single_phase_touching_X': 0.21,
                           'single_phase_spaced_R': 0.98, 'single_phase_spaced_X': 0.29,
                           'three_phase_trefoil_R': 0.86, 'three_phase_trefoil_X': 0.18,
                           'three_phase_flat_touching_R': 0.84, 'three_phase_flat_touching_X': 0.25,
                           'three_phase_flat_spaced_R': 0.84, 'three_phase_flat_spaced_X': 0.33}},
    70.0: {'voltage_drop': {'single_phase_touching_R': 0.68, 'single_phase_touching_X': 0.21,
                           'single_phase_spaced_R': 0.69, 'single_phase_spaced_X': 0.29,
                           'three_phase_trefoil_R': 0.59, 'three_phase_trefoil_X': 0.17,
                           'three_phase_flat_touching_R': 0.60, 'three_phase_flat_touching_X': 0.25,
                           'three_phase_flat_spaced_R': 0.62, 'three_phase_flat_spaced_X': 0.32}},
    95.0: {'voltage_drop': {'single_phase_touching_R': 0.51, 'single_phase_touching_X': 0.195,
                           'single_phase_spaced_R': 0.53, 'single_phase_spaced_X': 0.28,
                           'three_phase_trefoil_R': 0.44, 'three_phase_trefoil_X': 0.17,
                           'three_phase_flat_touching_R': 0.46, 'three_phase_flat_touching_X': 0.24,
                           'three_phase_flat_spaced_R': 0.49, 'three_phase_flat_spaced_X': 0.31}},
    120.0: {'voltage_drop': {'single_phase_touching_R': 0.41, 'single_phase_touching_X': 0.19,
                            'single_phase_spaced_R': 0.43, 'single_phase_spaced_X': 0.27,
                            'three_phase_trefoil_R': 0.35, 'three_phase_trefoil_X': 0.165,
                            'three_phase_flat_touching_R': 0.38, 'three_phase_flat_touching_X': 0.24,
                            'three_phase_flat_spaced_R': 0.41, 'three_phase_flat_spaced_X': 0.30}},
    150.0: {'voltage_drop': {'single_phase_touching_R': 0.33, 'single_phase_touching_X': 0.185,
                            'single_phase_spaced_R': 0.36, 'single_phase_spaced_X': 0.27,
                            'three_phase_trefoil_R': 0.29, 'three_phase_trefoil_X': 0.16,
                            'three_phase_flat_touching_R': 0.31, 'three_phase_flat_touching_X': 0.23,
                            'three_phase_flat_spaced_R': 0.34, 'three_phase_flat_spaced_X': 0.29}},
    185.0: {'voltage_drop': {'single_phase_touching_R': 0.27, 'single_phase_touching_X': 0.185,
                            'single_phase_spaced_R': 0.30, 'single_phase_spaced_X': 0.26,
                            'three_phase_trefoil_R': 0.23, 'three_phase_trefoil_X': 0.16,
                            'three_phase_flat_touching_R': 0.26, 'three_phase_flat_touching_X': 0.23,
                            'three_phase_flat_spaced_R': 0.29, 'three_phase_flat_spaced_X': 0.29}},
    240.0: {'voltage_drop': {'single_phase_touching_R': 0.21, 'single_phase_touching_X': 0.18,
                            'single_phase_spaced_R': 0.24, 'single_phase_spaced_X': 0.26,
                            'three_phase_trefoil_R': 0.18, 'three_phase_trefoil_X': 0.155,
                            'three_phase_flat_touching_R': 0.21, 'three_phase_flat_touching_X': 0.22,
                            'three_phase_flat_spaced_R': 0.24, 'three_phase_flat_spaced_X': 0.28}},
    300.0: {'voltage_drop': {'single_phase_touching_R': 0.17, 'single_phase_touching_X': 0.175,
                            'single_phase_spaced_R': 0.195, 'single_phase_spaced_X': 0.25,
                            'three_phase_trefoil_R': 0.145, 'three_phase_trefoil_X': 0.15,
                            'three_phase_flat_touching_R': 0.17, 'three_phase_flat_touching_X': 0.22,
                            'three_phase_flat_spaced_R': 0.20, 'three_phase_flat_spaced_X': 0.27}},
    400.0: {'voltage_drop': {'single_phase_touching_R': 0.145, 'single_phase_touching_X': 0.17,
                            'single_phase_spaced_R': 0.18, 'single_phase_spaced_X': 0.24,
                            'three_phase_trefoil_R': 0.125, 'three_phase_trefoil_X': 0.15,
                            'three_phase_flat_touching_R': 0.16, 'three_phase_flat_touching_X': 0.21,
                            'three_phase_flat_spaced_R': 0.20, 'three_phase_flat_spaced_X': 0.27}},
    500.0: {'voltage_drop': {'single_phase_touching_R': 0.125, 'single_phase_touching_X': 0.17,
                            'single_phase_spaced_R': 0.165, 'single_phase_spaced_X': 0.24,
                            'three_phase_trefoil_R': 0.105, 'three_phase_trefoil_X': 0.145,
                            'three_phase_flat_touching_R': 0.145, 'three_phase_flat_touching_X': 0.20,
                            'three_phase_flat_spaced_R': 0.19, 'three_phase_flat_spaced_X': 0.24}},
    630.0: {'voltage_drop': {'single_phase_touching_R': 0.105, 'single_phase_touching_X': 0.165,
                            'single_phase_spaced_R': 0.15, 'single_phase_spaced_X': 0.23,
                            'three_phase_trefoil_R': 0.092, 'three_phase_trefoil_X': 0.145,
                            'three_phase_flat_touching_R': 0.135, 'three_phase_flat_touching_X': 0.195,
                            'three_phase_flat_spaced_R': 0.175, 'three_phase_flat_spaced_X': 0.23}},
    800.0: {'voltage_drop': {'single_phase_touching_R': 0.09, 'single_phase_touching_X': 0.16,
                            'single_phase_spaced_R': 0.145, 'single_phase_spaced_X': 0.23,
                            'three_phase_trefoil_R': 0.086, 'three_phase_trefoil_X': 0.14,
                            'three_phase_flat_touching_R': 0.13, 'three_phase_flat_touching_X': 0.18,
                            'three_phase_flat_spaced_R': 0.175, 'three_phase_flat_spaced_X': 0.195}},
    1000.0: {'voltage_drop': {'single_phase_touching_R': 0.092, 'single_phase_touching_X': 0.155,
                             'single_phase_spaced_R': 0.14, 'single_phase_spaced_X': 0.21,
                             'three_phase_trefoil_R': 0.08, 'three_phase_trefoil_X': 0.135,
                             'three_phase_flat_touching_R': 0.125, 'three_phase_flat_touching_X': 0.17,
                             'three_phase_flat_spaced_R': 0.165, 'three_phase_flat_spaced_X': 0.18}},
}

# MULTI CORE ARMOURED (Table 4E4B)
XLPE_90_MULTI_ARMOURED = {
    1.5: {'voltage_drop': {'single_phase_mv': 31, 'three_phase_mv': 27}},
    2.5: {'voltage_drop': {'single_phase_mv': 19, 'three_phase_mv': 16}},
    4.0: {'voltage_drop': {'single_phase_mv': 12, 'three_phase_mv': 10}},
    6.0: {'voltage_drop': {'single_phase_mv': 7.9, 'three_phase_mv': 6.8}},
    10.0: {'voltage_drop': {'single_phase_mv': 4.7, 'three_phase_mv': 4.0}},
    16.0: {'voltage_drop': {'single_phase_mv': 2.9, 'three_phase_mv': 2.5}},
    25.0: {'voltage_drop': {'single_phase_R': 1.85, 'single_phase_X': 0.16, 'three_phase_R': 1.6, 'three_phase_X': 0.14}},
    35.0: {'voltage_drop': {'single_phase_R': 1.35, 'single_phase_X': 0.155, 'three_phase_R': 1.15, 'three_phase_X': 0.135}},
    50.0: {'voltage_drop': {'single_phase_R': 0.99, 'single_phase_X': 0.155, 'three_phase_R': 0.86, 'three_phase_X': 0.135}},
    70.0: {'voltage_drop': {'single_phase_R': 0.67, 'single_phase_X': 0.15, 'three_phase_R': 0.59, 'three_phase_X': 0.13}},
    95.0: {'voltage_drop': {'single_phase_R': 0.50, 'single_phase_X': 0.15, 'three_phase_R': 0.43, 'three_phase_X': 0.13}},
    120.0: {'voltage_drop': {'single_phase_R': 0.40, 'single_phase_X': 0.145, 'three_phase_R': 0.34, 'three_phase_X': 0.13}},
    150.0: {'voltage_drop': {'single_phase_R': 0.32, 'single_phase_X': 0.145, 'three_phase_R': 0.28, 'three_phase_X': 0.125}},
    185.0: {'voltage_drop': {'single_phase_R': 0.26, 'single_phase_X': 0.145, 'three_phase_R': 0.22, 'three_phase_X': 0.125}},
    240.0: {'voltage_drop': {'single_phase_R': 0.20, 'single_phase_X': 0.14, 'three_phase_R': 0.175, 'three_phase_X': 0.125}},
    300.0: {'voltage_drop': {'single_phase_R': 0.16, 'single_phase_X': 0.14, 'three_phase_R': 0.14, 'three_phase_X': 0.12}},
    400.0: {'voltage_drop': {'single_phase_R': 0.13, 'single_phase_X': 0.14, 'three_phase_R': 0.115, 'three_phase_X': 0.12}},
}

# ========== VOLTAGE DROP LOOKUP FUNCTION ==========
def get_voltage_drop_values(cable_type, size_mm2, phase, formation='flat'):
    """Get accurate voltage drop values based on cable type, size, phase, and formation"""
    
    if cable_type == 'multi_core_non_armoured':
        if size_mm2 <= 16:
            data = XLPE_90_MULTI_NON_ARMOURED.get(size_mm2, {})
            vd = data.get('voltage_drop', {})
            if phase == '3-phase':
                return {'type': 'mv', 'value': vd.get('three_phase_mv', 0)}
            else:
                return {'type': 'mv', 'value': vd.get('single_phase_mv', vd.get('dc_mv', 0))}
        else:
            data = XLPE_90_MULTI_NON_ARMOURED.get(size_mm2, {})
            vd = data.get('voltage_drop', {})
            if phase == '3-phase':
                return {'type': 'rx', 'R': vd.get('three_phase_R', 0), 'X': vd.get('three_phase_X', 0)}
            else:
                return {'type': 'rx', 'R': vd.get('single_phase_R', 0), 'X': vd.get('single_phase_X', 0)}
    
    elif cable_type == 'single_core_non_armoured':
        if size_mm2 <= 16:
            data = XLPE_90_SINGLE_NON_ARMOURED.get(size_mm2, {})
            vd = data.get('voltage_drop', {})
            if phase == '3-phase':
                return {'type': 'mv', 'value': vd.get('three_phase_mv', 0)}
            else:
                return {'type': 'mv', 'value': vd.get('single_phase_mv', 0)}
        else:
            data = XLPE_90_SINGLE_NON_ARMOURED.get(size_mm2, {})
            vd = data.get('voltage_drop', {})
            if phase == '3-phase':
                return {'type': 'rx', 'R': vd.get('three_phase_R', 0), 'X': vd.get('three_phase_X', 0)}
            else:
                return {'type': 'rx', 'R': vd.get('single_phase_R', 0), 'X': vd.get('single_phase_X', 0)}
    
    elif cable_type == 'single_core_armoured':
        data = XLPE_90_SINGLE_ARMOURED.get(size_mm2, {})
        vd = data.get('voltage_drop', {})
        
        if phase == '3-phase':
            if formation == 'trefoil':
                return {'type': 'rx', 'R': vd.get('three_phase_trefoil_R', 0), 'X': vd.get('three_phase_trefoil_X', 0)}
            elif formation == 'spaced':
                return {'type': 'rx', 'R': vd.get('three_phase_flat_spaced_R', 0), 'X': vd.get('three_phase_flat_spaced_X', 0)}
            else:
                return {'type': 'rx', 'R': vd.get('three_phase_flat_touching_R', 0), 'X': vd.get('three_phase_flat_touching_X', 0)}
        else:
            if formation == 'spaced':
                return {'type': 'rx', 'R': vd.get('single_phase_spaced_R', 0), 'X': vd.get('single_phase_spaced_X', 0)}
            else:
                return {'type': 'rx', 'R': vd.get('single_phase_touching_R', 0), 'X': vd.get('single_phase_touching_X', 0)}
    
    elif cable_type == 'multi_core_armoured':
        if size_mm2 <= 16:
            data = XLPE_90_MULTI_ARMOURED.get(size_mm2, {})
            vd = data.get('voltage_drop', {})
            if phase == '3-phase':
                return {'type': 'mv', 'value': vd.get('three_phase_mv', 0)}
            else:
                return {'type': 'mv', 'value': vd.get('single_phase_mv', 0)}
        else:
            data = XLPE_90_MULTI_ARMOURED.get(size_mm2, {})
            vd = data.get('voltage_drop', {})
            if phase == '3-phase':
                return {'type': 'rx', 'R': vd.get('three_phase_R', 0), 'X': vd.get('three_phase_X', 0)}
            else:
                return {'type': 'rx', 'R': vd.get('single_phase_R', 0), 'X': vd.get('single_phase_X', 0)}
    
    return {'type': 'mv', 'value': 0}

# ========== TABLE CONFIGURATION FUNCTIONS ==========
def get_valid_reference_methods(cable_type):
    valid_methods = {
        'single_core_non_armoured': ['B', 'C', 'F', 'G'],
        'multi_core_non_armoured': ['B', 'C', 'E'],
        'single_core_armoured': ['C', 'F'],
        'multi_core_armoured': ['C', 'E', 'D', 'D_direct']
    }
    return valid_methods.get(cable_type, ['B', 'C'])

def get_table_configurations(cable_type, reference_method):
    configs = {
        'single_core_non_armoured': {
            'B': [{'key': 'B2', 'description': '2 cables, single-phase a.c. or d.c.', 'phase': '1-phase'},
                  {'key': 'B34', 'description': '3 or 4 cables, three-phase a.c.', 'phase': '3-phase'}],
            'C': [{'key': 'C2', 'description': '2 cables, single-phase a.c. or d.c.', 'phase': '1-phase'},
                  {'key': 'C34', 'description': '3 or 4 cables, three-phase a.c.', 'phase': '3-phase'}],
            'F': [{'key': 'F2', 'description': '2 cables, flat and touching', 'phase': '1-phase'},
                  {'key': 'F34_flat', 'description': '3 or 4 cables, flat and touching', 'phase': '3-phase'},
                  {'key': 'F34_trefoil', 'description': '3 cables, trefoil formation', 'phase': '3-phase'}],
            'G': [{'key': 'G2', 'description': '2 cables, spaced', 'phase': '1-phase'},
                  {'key': 'G34', 'description': '3 cables, spaced', 'phase': '3-phase'}]
        },
        'multi_core_non_armoured': {
            'B': [{'key': 'B2', 'description': '1 two-core cable', 'phase': '1-phase'},
                  {'key': 'B34', 'description': '1 three/four-core cable', 'phase': '3-phase'}],
            'C': [{'key': 'C2', 'description': '1 two-core cable', 'phase': '1-phase'},
                  {'key': 'C34', 'description': '1 three/four-core cable', 'phase': '3-phase'}],
            'E': [{'key': 'E2', 'description': '1 two-core cable', 'phase': '1-phase'},
                  {'key': 'E34', 'description': '1 three/four-core cable', 'phase': '3-phase'}]
        },
        'single_core_armoured': {
            'C': [{'key': 'C2', 'description': '2 cables, flat and touching', 'phase': '1-phase'},
                  {'key': 'C34', 'description': '3 or 4 cables, flat and touching', 'phase': '3-phase'}],
            'F': [{'key': 'F2_flat', 'description': '2 cables, flat and touching', 'phase': '1-phase'},
                  {'key': 'F34_flat', 'description': '3 cables, flat and touching', 'phase': '3-phase'},
                  {'key': 'F34_trefoil', 'description': '3 cables, trefoil formation', 'phase': '3-phase'}]
        },
        'multi_core_armoured': {
            'C': [{'key': 'C2', 'description': '1 two-core cable', 'phase': '1-phase'},
                  {'key': 'C34', 'description': '1 three/four-core cable', 'phase': '3-phase'}],
            'E': [{'key': 'E2', 'description': '1 two-core cable', 'phase': '1-phase'},
                  {'key': 'E34', 'description': '1 three/four-core cable', 'phase': '3-phase'}],
            'D': [{'key': 'D2', 'description': '1 two-core cable (buried)', 'phase': '1-phase'},
                  {'key': 'D34', 'description': '1 three/four-core cable (buried)', 'phase': '3-phase'}],
            'D_direct': [{'key': 'D2', 'description': '1 two-core cable (direct buried)', 'phase': '1-phase'},
                         {'key': 'D34', 'description': '1 three/four-core cable (direct buried)', 'phase': '3-phase'}]
        }
    }
    return configs.get(cable_type, {}).get(reference_method, [])

def get_table_reference_info(cable_type, reference_method, config_key):
    table_info = {
        'single_core_non_armoured': {'B': 'BS 7671 Table 4E1A', 'C': 'BS 7671 Table 4E1A', 'F': 'BS 7671 Table 4E1A', 'G': 'BS 7671 Table 4E1A'},
        'multi_core_non_armoured': {'B': 'BS 7671 Table 4E2A', 'C': 'BS 7671 Table 4E2A', 'E': 'BS 7671 Table 4E2A'},
        'single_core_armoured': {'C': 'BS 7671 Table 4E3A', 'F': 'BS 7671 Table 4E3A'},
        'multi_core_armoured': {'C': 'BS 7671 Table 4E4A', 'E': 'BS 7671 Table 4E4A', 'D': 'BS 7671 Table 4E4A', 'D_direct': 'BS 7671 Table 4E4A'}
    }
    return table_info.get(cable_type, {}).get(reference_method, 'IEC 60502-2 / BS 7671')

def get_valid_arrangements(installation_method, cable_type):
    if installation_method in ['D', 'D_direct']:
        if 'multi_core_armoured' in cable_type:
            return ['direct_buried', 'buried_ducts']
        else:
            return ['bunched_in_air_surface_enclosed', 'single_layer_wall_floor', 
                    'single_layer_perforated_tray', 'single_layer_ladder_cleats']
    else:
        if installation_method == 'B':
            return ['bunched_in_air_surface_enclosed']
        elif installation_method == 'C':
            return ['bunched_in_air_surface_enclosed', 'single_layer_wall_floor']
        elif installation_method == 'E':
            return ['single_layer_perforated_tray']
        elif installation_method == 'F':
            return ['single_layer_ladder_cleats']
        elif installation_method == 'G':
            return ['single_layer_ladder_cleats']
        else:
            return ['bunched_in_air_surface_enclosed', 'single_layer_wall_floor', 
                    'single_layer_perforated_tray', 'single_layer_ladder_cleats']

def get_clearance_options(installation_method, arrangement, is_single_core=True):
    if installation_method in ['D', 'D_direct']:
        if arrangement == 'direct_buried':
            return ['touching', 'one_diameter', 'clearance_0_125m', 'clearance_0_25m', 'clearance_0_5m']
        elif arrangement == 'buried_ducts':
            return ['touching', 'clearance_0_25m', 'clearance_0_5m', 'clearance_1_0m']
    return []

def get_ampacity_from_config(cable_data, table_config, load_phase, installation_method='C'):
    """Get base ampacity from table configuration based on installation method"""
    # First try exact match with table_config
    if table_config and table_config in cable_data.get('ampacity', {}):
        return cable_data['ampacity'][table_config]
    
    # If not found, use method-appropriate fallback
    if installation_method in ['D', 'D_direct']:
        # For buried installations - use D2/D34 keys from Table 4E4A
        if load_phase == '3-phase':
            return cable_data.get('D34', cable_data.get('C34', cable_data.get('E34', 0)))
        else:
            return cable_data.get('D2', cable_data.get('C2', cable_data.get('E2', 0)))
    elif installation_method in ['E']:
        # For perforated tray - use E2/E34 keys
        if load_phase == '3-phase':
            return cable_data.get('E34', cable_data.get('C34', 0))
        else:
            return cable_data.get('E2', cable_data.get('C2', 0))
    elif installation_method in ['F', 'G']:
        # For free air - use F keys
        if load_phase == '3-phase':
            return cable_data.get('F34_flat', cable_data.get('F34_trefoil', cable_data.get('C34', 0)))
        else:
            return cable_data.get('F2', cable_data.get('C2', 0))
    else:
        # Default to Method C
        if load_phase == '3-phase':
            return cable_data.get('C34', cable_data.get('E34', 0))
        else:
            return cable_data.get('C2', cable_data.get('E2', 0))

# ========== CABLE DATABASE WITH AMPACITIES ==========
# Multi-core non-armoured ampacities (Table 4E2A)
XLPE_90_MULTI_NON_ARMOURED_AMP = {
    1.0: {'B2': 17.0, 'B34': 15.0, 'C2': 19.0, 'C34': 17.0, 'E2': 21.0, 'E34': 18.0},
    1.5: {'B2': 22.0, 'B34': 19.5, 'C2': 24.0, 'C34': 22.0, 'E2': 26.0, 'E34': 23.0},
    2.5: {'B2': 30.0, 'B34': 26.0, 'C2': 33.0, 'C34': 30.0, 'E2': 36.0, 'E34': 32.0},
    4.0: {'B2': 40.0, 'B34': 35.0, 'C2': 45.0, 'C34': 40.0, 'E2': 49.0, 'E34': 42.0},
    6.0: {'B2': 51.0, 'B34': 44.0, 'C2': 58.0, 'C34': 52.0, 'E2': 63.0, 'E34': 54.0},
    10.0: {'B2': 69.0, 'B34': 60.0, 'C2': 80.0, 'C34': 71.0, 'E2': 86.0, 'E34': 75.0},
    16.0: {'B2': 91.0, 'B34': 80.0, 'C2': 107.0, 'C34': 96.0, 'E2': 115.0, 'E34': 100.0},
    25.0: {'B2': 119.0, 'B34': 105.0, 'C2': 138.0, 'C34': 119.0, 'E2': 149.0, 'E34': 127.0},
    35.0: {'B2': 146.0, 'B34': 128.0, 'C2': 171.0, 'C34': 147.0, 'E2': 185.0, 'E34': 158.0},
    50.0: {'B2': 175.0, 'B34': 154.0, 'C2': 209.0, 'C34': 179.0, 'E2': 225.0, 'E34': 192.0},
    70.0: {'B2': 221.0, 'B34': 194.0, 'C2': 269.0, 'C34': 229.0, 'E2': 289.0, 'E34': 246.0},
    95.0: {'B2': 265.0, 'B34': 233.0, 'C2': 328.0, 'C34': 278.0, 'E2': 352.0, 'E34': 298.0},
    120.0: {'B2': 305.0, 'B34': 268.0, 'C2': 382.0, 'C34': 322.0, 'E2': 410.0, 'E34': 346.0},
    150.0: {'B2': 334.0, 'B34': 300.0, 'C2': 441.0, 'C34': 371.0, 'E2': 473.0, 'E34': 399.0},
    185.0: {'B2': 384.0, 'B34': 340.0, 'C2': 506.0, 'C34': 424.0, 'E2': 542.0, 'E34': 456.0},
    240.0: {'B2': 459.0, 'B34': 398.0, 'C2': 599.0, 'C34': 500.0, 'E2': 641.0, 'E34': 538.0},
    300.0: {'B2': 532.0, 'B34': 455.0, 'C2': 693.0, 'C34': 576.0, 'E2': 741.0, 'E34': 621.0},
    400.0: {'B2': 625.0, 'B34': 536.0, 'C2': 803.0, 'C34': 667.0, 'E2': 865.0, 'E34': 741.0},
}

# Single core non-armoured ampacities (Table 4E1A)
XLPE_90_SINGLE_NON_ARMOURED_AMP = {
    1.5: {'B2': 23.0, 'B34': 20.0, 'C2': 25.0, 'C34': 23.0, 'F2': 26.0, 'F34_flat': 23.0, 'F34_trefoil': 23.0, 'G2': 32.0, 'G34': 29.0},
    2.5: {'B2': 31.0, 'B34': 28.0, 'C2': 34.0, 'C34': 31.0, 'F2': 36.0, 'F34_flat': 32.0, 'F34_trefoil': 32.0, 'G2': 44.0, 'G34': 40.0},
    4.0: {'B2': 42.0, 'B34': 37.0, 'C2': 46.0, 'C34': 41.0, 'F2': 49.0, 'F34_flat': 42.0, 'F34_trefoil': 42.0, 'G2': 59.0, 'G34': 53.0},
    6.0: {'B2': 54.0, 'B34': 48.0, 'C2': 59.0, 'C34': 54.0, 'F2': 63.0, 'F34_flat': 54.0, 'F34_trefoil': 54.0, 'G2': 76.0, 'G34': 68.0},
    10.0: {'B2': 75.0, 'B34': 66.0, 'C2': 81.0, 'C34': 74.0, 'F2': 86.0, 'F34_flat': 75.0, 'F34_trefoil': 75.0, 'G2': 103.0, 'G34': 93.0},
    16.0: {'B2': 100.0, 'B34': 88.0, 'C2': 109.0, 'C34': 99.0, 'F2': 115.0, 'F34_flat': 100.0, 'F34_trefoil': 100.0, 'G2': 138.0, 'G34': 124.0},
    25.0: {'B2': 133.0, 'B34': 117.0, 'C2': 143.0, 'C34': 130.0, 'F2': 149.0, 'F34_flat': 127.0, 'F34_trefoil': 127.0, 'G2': 176.0, 'G34': 158.0},
    35.0: {'B2': 164.0, 'B34': 144.0, 'C2': 176.0, 'C34': 161.0, 'F2': 185.0, 'F34_flat': 158.0, 'F34_trefoil': 158.0, 'G2': 218.0, 'G34': 196.0},
    50.0: {'B2': 198.0, 'B34': 175.0, 'C2': 228.0, 'C34': 209.0, 'F2': 225.0, 'F34_flat': 192.0, 'F34_trefoil': 192.0, 'G2': 265.0, 'G34': 240.0},
    70.0: {'B2': 253.0, 'B34': 222.0, 'C2': 293.0, 'C34': 268.0, 'F2': 289.0, 'F34_flat': 246.0, 'F34_trefoil': 246.0, 'G2': 340.0, 'G34': 307.0},
    95.0: {'B2': 306.0, 'B34': 269.0, 'C2': 355.0, 'C34': 326.0, 'F2': 352.0, 'F34_flat': 298.0, 'F34_trefoil': 298.0, 'G2': 415.0, 'G34': 375.0},
    120.0: {'B2': 354.0, 'B34': 312.0, 'C2': 413.0, 'C34': 379.0, 'F2': 410.0, 'F34_flat': 346.0, 'F34_trefoil': 346.0, 'G2': 483.0, 'G34': 437.0},
    150.0: {'B2': 393.0, 'B34': 342.0, 'C2': 476.0, 'C34': 436.0, 'F2': 473.0, 'F34_flat': 399.0, 'F34_trefoil': 399.0, 'G2': 558.0, 'G34': 505.0},
    185.0: {'B2': 449.0, 'B34': 384.0, 'C2': 545.0, 'C34': 500.0, 'F2': 542.0, 'F34_flat': 456.0, 'F34_trefoil': 456.0, 'G2': 640.0, 'G34': 580.0},
    240.0: {'B2': 528.0, 'B34': 450.0, 'C2': 644.0, 'C34': 590.0, 'F2': 641.0, 'F34_flat': 538.0, 'F34_trefoil': 538.0, 'G2': 757.0, 'G34': 686.0},
    300.0: {'B2': 603.0, 'B34': 514.0, 'C2': 743.0, 'C34': 681.0, 'F2': 741.0, 'F34_flat': 621.0, 'F34_trefoil': 621.0, 'G2': 875.0, 'G34': 793.0},
    400.0: {'B2': 683.0, 'B34': 584.0, 'C2': 868.0, 'C34': 793.0, 'F2': 865.0, 'F34_flat': 741.0, 'F34_trefoil': 741.0, 'G2': 1025.0, 'G34': 930.0},
}

# Single core armoured ampacities (Table 4E3A)
XLPE_90_SINGLE_ARMOURED_AMP = {
    50.0: {'C2': 237.0, 'C34': 220.0, 'F2_flat': 253.0, 'F34_flat': 232.0, 'F34_trefoil': 222.0},
    70.0: {'C2': 303.0, 'C34': 277.0, 'F2_flat': 322.0, 'F34_flat': 293.0, 'F34_trefoil': 285.0},
    95.0: {'C2': 367.0, 'C34': 333.0, 'F2_flat': 389.0, 'F34_flat': 352.0, 'F34_trefoil': 346.0},
    120.0: {'C2': 425.0, 'C34': 383.0, 'F2_flat': 449.0, 'F34_flat': 405.0, 'F34_trefoil': 402.0},
    150.0: {'C2': 488.0, 'C34': 437.0, 'F2_flat': 516.0, 'F34_flat': 462.0, 'F34_trefoil': 463.0},
    185.0: {'C2': 557.0, 'C34': 496.0, 'F2_flat': 587.0, 'F34_flat': 524.0, 'F34_trefoil': 529.0},
    240.0: {'C2': 656.0, 'C34': 579.0, 'F2_flat': 689.0, 'F34_flat': 612.0, 'F34_trefoil': 625.0},
    300.0: {'C2': 755.0, 'C34': 662.0, 'F2_flat': 792.0, 'F34_flat': 700.0, 'F34_trefoil': 720.0},
    400.0: {'C2': 853.0, 'C34': 717.0, 'F2_flat': 899.0, 'F34_flat': 767.0, 'F34_trefoil': 815.0},
    500.0: {'C2': 962.0, 'C34': 791.0, 'F2_flat': 1016.0, 'F34_flat': 851.0, 'F34_trefoil': 918.0},
    630.0: {'C2': 1082.0, 'C34': 861.0, 'F2_flat': 1146.0, 'F34_flat': 935.0, 'F34_trefoil': 1027.0},
    800.0: {'C2': 1170.0, 'C34': 904.0, 'F2_flat': 1246.0, 'F34_flat': 987.0, 'F34_trefoil': 1119.0},
    1000.0: {'C2': 1261.0, 'C34': 961.0, 'F2_flat': 1345.0, 'F34_flat': 1055.0, 'F34_trefoil': 1214.0},
}

# Multi-core armoured ampacities (Table 4E4A)
XLPE_90_MULTI_ARMOURED_AMP = {
    1.5: {'C2': 27.0, 'C34': 23.0, 'E2': 29.0, 'E34': 25.0, 'D2': 25.0, 'D34': 21.0},
    2.5: {'C2': 36.0, 'C34': 31.0, 'E2': 39.0, 'E34': 33.0, 'D2': 33.0, 'D34': 28.0},
    4.0: {'C2': 49.0, 'C34': 42.0, 'E2': 52.0, 'E34': 44.0, 'D2': 43.0, 'D34': 36.0},
    6.0: {'C2': 62.0, 'C34': 53.0, 'E2': 66.0, 'E34': 56.0, 'D2': 53.0, 'D34': 44.0},
    10.0: {'C2': 85.0, 'C34': 73.0, 'E2': 90.0, 'E34': 78.0, 'D2': 71.0, 'D34': 58.0},
    16.0: {'C2': 110.0, 'C34': 94.0, 'E2': 115.0, 'E34': 99.0, 'D2': 91.0, 'D34': 75.0},
    25.0: {'C2': 146.0, 'C34': 124.0, 'E2': 152.0, 'E34': 131.0, 'D2': 116.0, 'D34': 96.0},
    35.0: {'C2': 180.0, 'C34': 154.0, 'E2': 188.0, 'E34': 162.0, 'D2': 139.0, 'D34': 115.0},
    50.0: {'C2': 219.0, 'C34': 187.0, 'E2': 228.0, 'E34': 197.0, 'D2': 164.0, 'D34': 135.0},
    70.0: {'C2': 279.0, 'C34': 238.0, 'E2': 291.0, 'E34': 251.0, 'D2': 203.0, 'D34': 167.0},
    95.0: {'C2': 338.0, 'C34': 289.0, 'E2': 354.0, 'E34': 304.0, 'D2': 239.0, 'D34': 197.0},
    120.0: {'C2': 392.0, 'C34': 335.0, 'E2': 410.0, 'E34': 353.0, 'D2': 271.0, 'D34': 223.0},
    150.0: {'C2': 451.0, 'C34': 386.0, 'E2': 472.0, 'E34': 406.0, 'D2': 306.0, 'D34': 251.0},
    185.0: {'C2': 515.0, 'C34': 441.0, 'E2': 539.0, 'E34': 463.0, 'D2': 343.0, 'D34': 281.0},
    240.0: {'C2': 607.0, 'C34': 520.0, 'E2': 636.0, 'E34': 546.0, 'D2': 395.0, 'D34': 324.0},
    300.0: {'C2': 698.0, 'C34': 599.0, 'E2': 732.0, 'E34': 628.0, 'D2': 446.0, 'D34': 365.0},
    400.0: {'C2': 787.0, 'C34': 673.0, 'E2': 847.0, 'E34': 728.0},
}

def get_cable_ampacities(cable_type):
    if cable_type == 'multi_core_non_armoured':
        return XLPE_90_MULTI_NON_ARMOURED_AMP
    elif cable_type == 'single_core_non_armoured':
        return XLPE_90_SINGLE_NON_ARMOURED_AMP
    elif cable_type == 'single_core_armoured':
        return XLPE_90_SINGLE_ARMOURED_AMP
    elif cable_type == 'multi_core_armoured':
        return XLPE_90_MULTI_ARMOURED_AMP
    return {}

# ========== CABLE SIZING CALCULATOR CLASS ==========
class CableSizingCalculator:
    def __init__(self):
        self.results = {}
    
    def calculate_load_current(self, power_kw, voltage_v, pf, efficiency=1.0, phase='3-phase'):
        if phase == '3-phase':
            return (power_kw * 1000) / (1.732 * voltage_v * pf * efficiency)
        elif phase == '1-phase':
            return (power_kw * 1000) / (voltage_v * pf * efficiency)
        else:
            return (power_kw * 1000) / voltage_v
    
    def calculate_operating_temperature(self, ambient_temp, load_current, derated_ampacity, insulation_temp):
        if derated_ampacity > 0:
            load_ratio = load_current / derated_ampacity
            temp_rise = (insulation_temp - ambient_temp) * (load_ratio ** 2)
            operating_temp = ambient_temp + temp_rise
        else:
            operating_temp = insulation_temp
        return operating_temp
    
    def calculate_short_circuit(self, size_mm2, insulation_type, ambient_temp, load_current, rated_current, k1, k2, k3, k4, duration_s=1.0, conductor_material='Copper'):
        Isc, K, θi, θf = calculate_short_circuit_current(size_mm2, insulation_type, duration_s, conductor_material)
        insulation_temp = 70 if insulation_type == 'PVC' else 90
        total_k = k1 * k2 * k3 * k4
        derated_ampacity = rated_current * total_k
        operating_temp = self.calculate_operating_temperature(ambient_temp, load_current, derated_ampacity, insulation_temp)
        return Isc, K, operating_temp, θi, θf
    
    def get_derating_factors(self, temp_c, insulation_temp, num_cables, arrangement, installation, 
                             soil_resistivity, depth, is_single_core=True, clearance='touching'):
        if installation in ['buried', 'duct', 'trench', 'ground', 'D', 'D_direct']:
            k1 = get_temperature_factor(insulation_temp, temp_c, 'ground')
            install_type = 'buried'
        else:
            k1 = get_temperature_factor(insulation_temp, temp_c, 'air')
            install_type = 'air'
        
        k2 = get_grouping_factor(num_cables, arrangement, install_type, clearance, is_single_core)
        
        if installation in ['buried', 'duct', 'ground', 'D', 'D_direct']:
            k3 = get_soil_resistivity_factor(soil_resistivity, installation, is_single_core)
            k4 = get_depth_factor(depth, installation, is_single_core)
        else:
            k3 = 1.0
            k4 = 1.0
        
        total_k = k1 * k2 * k3 * k4
        factors = {'k1 (Temperature)': k1, 'k2 (Grouping)': k2, 'k3 (Soil Resistivity)': k3, 'k4 (Depth)': k4, 'total': total_k}
        return total_k, factors
    
    def calculate_voltage_drop(self, current, length_m, cable_type, size_mm2, pf, voltage_v, phase='3-phase', formation='flat'):
        """Calculate voltage drop using accurate values - NO derating factors applied"""
        vd_values = get_voltage_drop_values(cable_type, size_mm2, phase, formation)
        
        if vd_values['type'] == 'mv':
            mv_per_am = vd_values['value']
            if mv_per_am == 0:
                return 0, 0
            Vd = mv_per_am * current * length_m / 1000
            vd_percent = (Vd / voltage_v) * 100
            return Vd, vd_percent
        else:
            r = vd_values['R']
            x = vd_values['X']
            if r == 0 and x == 0:
                return 0, 0
            
            phi = math.acos(pf)
            sin_phi = math.sin(phi)
            
            if phase == '3-phase':
                Vd = 1.732 * current * (r * pf + x * sin_phi) * length_m / 1000
            else:
                Vd = 2 * current * (r * pf + x * sin_phi) * length_m / 1000
            
            vd_percent = (Vd / voltage_v) * 100
            return Vd, vd_percent
    
    def get_cable_category(self, voltage_v):
        if voltage_v <= 1000:
            return 'LV (0.6/1kV)', 'LV'
        elif voltage_v <= 3300:
            return 'MV (3.3kV)', 'MV_33KV'
        elif voltage_v <= 6600:
            return 'MV (6.6kV)', 'MV_66KV'
        else:
            return 'MV (11kV)', 'MV_11KV'

def select_cable_automatically(load, cable_calc, ambient_temp, insulation_temp, load_current, 
                                load_length, load_pf, load_voltage, load_phase,
                                installation_method, cable_formation, load_cable_type,
                                load_arrangement, load_soil_res, load_depth,
                                load_num_cables, table_config, clearance='touching'):
    
    if 'single_core' in load_cable_type and installation_method in ['D', 'D_direct']:
        return None, None, 0, 0, 0, 0, {}, False, []
    
    cable_ampacities = get_cable_ampacities(load_cable_type)
    available_sizes = sorted(cable_ampacities.keys())
    
    for size in available_sizes:
        cable_data = cable_ampacities[size]
        
        # Get ampacity using installation method
        ampacity = get_ampacity_from_config(cable_data, table_config, load_phase, installation_method)
        
        if ampacity == 0:
            continue
        
        is_single_core = (load_cable_type in ['single_core_non_armoured', 'single_core_armoured'])
        
        total_k, factors = cable_calc.get_derating_factors(
            ambient_temp, insulation_temp,
            load_num_cables, load_arrangement,
            installation_method,
            load_soil_res, load_depth,
            is_single_core, clearance
        )
        
        derated = ampacity * total_k
        ampacity_pass = derated >= load_current
        
        vd_v, vd_pct = cable_calc.calculate_voltage_drop(
            load_current, load_length, load_cable_type, size,
            load_pf, load_voltage, load_phase, cable_formation
        )
        
        vd_pass = vd_pct <= 2.5
        
        if ampacity_pass and vd_pass:
            return size, cable_data, ampacity, derated, vd_pct, total_k, factors, True, []
    
    if available_sizes:
        largest_size = max(available_sizes)
        largest_data = cable_ampacities[largest_size]
        
        ampacity = get_ampacity_from_config(largest_data, table_config, load_phase, installation_method)
        
        if ampacity == 0:
            ampacity = largest_data.get('C2', 0)
        
        is_single_core = (load_cable_type in ['single_core_non_armoured', 'single_core_armoured'])
        total_k, factors = cable_calc.get_derating_factors(
            ambient_temp, insulation_temp,
            load_num_cables, load_arrangement,
            installation_method,
            load_soil_res, load_depth,
            is_single_core, clearance
        )
        
        derated = ampacity * total_k
        vd_v, vd_pct = cable_calc.calculate_voltage_drop(
            load_current, load_length, load_cable_type, largest_size,
            load_pf, load_voltage, load_phase, cable_formation
        )
        
        return largest_size, largest_data, ampacity, derated, vd_pct, total_k, factors, False, []
    
    return None, None, 0, 0, 0, 0, {}, False, []

class CircuitBreakerCalculator:
    def get_standard_rating(self, current, design_factor=1.25):
        required = current * design_factor
        for rating in CB_RATINGS:
            if rating >= required:
                return rating, required
        return CB_RATINGS[-1], required
    
    def get_breaker_type(self, rating, voltage=400):
        if voltage >= 1000:
            return 'MV Circuit Breaker', 'IEC 62271-100'
        else:
            if rating <= 125:
                return 'MCB', 'IEC 60898'
            elif rating <= 1600:
                return 'MCCB', 'IEC 60947-2'
            else:
                return 'ACB', 'IEC 60947-2'
    
    def calculate_cb_size(self, loads_df, design_factor=1.25, manufacturer='Schneider Electric'):
        results = []
        detailed_reasons = []
        
        for idx, load in loads_df.iterrows():
            if load['Phase'] == '3-phase':
                current = load['Power (kW)'] * 1000 / (1.732 * load['Voltage (V)'] * load['Power Factor'])
                phase_desc = "Three-phase"
            elif load['Phase'] == '1-phase':
                current = load['Power (kW)'] * 1000 / (load['Voltage (V)'] * load['Power Factor'])
                phase_desc = "Single-phase"
            else:
                current = load['Power (kW)'] * 1000 / load['Voltage (V)']
                phase_desc = "DC"
            
            rating, required = self.get_standard_rating(current, design_factor)
            breaker_type, standard = self.get_breaker_type(rating, load['Voltage (V)'])
            series = MANUFACTURERS[manufacturer][breaker_type] if breaker_type in MANUFACTURERS[manufacturer] else 'Standard series'
            
            results.append({
                'Load': load['Load Name'],
                'Power (kW)': load['Power (kW)'],
                'Voltage (V)': load['Voltage (V)'],
                'Phase': load['Phase'],
                'Load Type': load.get('Load Type', 'Continuous'),
                'Current (A)': current,
                'Required CB (A)': required,
                'Selected CB (A)': rating,
                'Breaker Type': breaker_type,
                'Standard': standard,
                'Manufacturer': manufacturer,
                'Series': series,
                'Power Factor': load.get('Power Factor', 0.85)
            })
            
            detailed_reasons.append({
                'load_name': load['Load Name'],
                'phase_desc': phase_desc,
                'voltage': load['Voltage (V)'],
                'current': current,
                'required': required,
                'selected': rating,
                'breaker_type': breaker_type,
                'standard': standard,
                'design_factor': design_factor,
                'manufacturer': manufacturer,
                'series': series
            })
        
        return results, detailed_reasons
    
    def calculate_main_cb_by_voltage(self, loads_df, design_factor=1.25):
        voltage_groups = loads_df.groupby('Voltage (V)')
        
        results = {}
        detailed_reasons = {}
        
        for voltage, group in voltage_groups:
            total_power = group['Power (kW)'].sum()
            
            if len(group) > 1:
                avg_pf = sum(group['Power (kW)'] * group['Power Factor']) / total_power if total_power > 0 else 0.85
            else:
                avg_pf = group.iloc[0]['Power Factor']
            
            if voltage >= 1000:
                current = total_power * 1000 / (1.732 * voltage * avg_pf)
                system_type = "MV (Medium Voltage)"
                if voltage <= 3300:
                    voltage_range = "3.3kV"
                elif voltage <= 6600:
                    voltage_range = "6.6kV"
                else:
                    voltage_range = "11kV"
            else:
                current = total_power * 1000 / (1.732 * voltage * avg_pf)
                system_type = "LV (Low Voltage)"
                voltage_range = f"{int(voltage)}V"
            
            required = current * design_factor
            selected, _ = self.get_standard_rating(current, design_factor)
            breaker_type, standard = self.get_breaker_type(selected, voltage)
            
            manufacturer = 'Schneider Electric'
            series = MANUFACTURERS[manufacturer].get(breaker_type, 'Standard series') if breaker_type in MANUFACTURERS[manufacturer] else 'Standard series'
            
            results[voltage] = {
                'voltage': voltage,
                'voltage_range': voltage_range,
                'system_type': system_type,
                'total_power': total_power,
                'avg_pf': avg_pf,
                'current': current,
                'required_cb': required,
                'selected_cb': selected,
                'breaker_type': breaker_type,
                'standard': standard,
                'manufacturer': manufacturer,
                'series': series,
                'num_loads': len(group)
            }
            
            loads_list = "\n".join([f"  - {row['Load Name']}: {row['Power (kW)']:.1f} kW, PF={row['Power Factor']}" 
                                    for _, row in group.iterrows()])
            
            detailed_reasons[voltage] = f"""
MAIN CIRCUIT BREAKER DETAILED CALCULATION - {system_type} ({voltage_range})

================================================================================

Step 1: Load analysis for {voltage_range} system
--------------------------------------------------------------------------------
Voltage level: {voltage} V ({system_type})
Number of loads in this group: {len(group)}

Loads in this voltage group:
{loads_list}

Total connected load: {total_power:.2f} kW
Weighted average power factor: {avg_pf:.3f}

Step 2: Total current calculation
--------------------------------------------------------------------------------
Formula: I = P x 1000 / (1.732 x V x PF)
I = {total_power:.2f} x 1000 / (1.732 x {voltage} x {avg_pf:.3f})
I = {current:.2f} A

Step 3: Circuit breaker sizing
--------------------------------------------------------------------------------
Safety factor: {design_factor} (25% safety margin for continuous loads)
Required rating = Required current × Safety factor
Required = {current:.2f} × {design_factor} = {required:.2f} A

Step 4: Standard rating selection
--------------------------------------------------------------------------------
Standard circuit breaker ratings (A): 6, 10, 16, 20, 25, 32, 40, 50, 63, 80, 100, 
125, 160, 200, 250, 315, 400, 500, 630, 800, 1000, 1250, 1600

Selected rating: {selected} A (next standard rating ≥ {required:.2f} A)

Step 5: Breaker type selection
--------------------------------------------------------------------------------
Based on:
- Voltage level: {voltage} V ({system_type})
- Rated current: {selected} A

Selected breaker type: {breaker_type}
Standard: {standard}

================================================================================
FINAL SELECTION: {selected} A {breaker_type} for {voltage_range} System
================================================================================
"""
        
        return results, detailed_reasons

# ========== WORD REPORT CLASSES ==========
class LightningWordReport:
    def __init__(self):
        self.doc = Document()
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
    
    def add_calculations(self, results, inputs):
        title = self.doc.add_heading('LIGHTNING PROTECTION CALCULATIONS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(f'Date: {format_pakistan_datetime()} (Pakistan Time)')
        self.doc.add_paragraph('Reference: IEC 62305-2')
        self.doc.add_paragraph()
        
        self.doc.add_heading('1.1 Collection area (Ad)', level=1)
        self.doc.add_paragraph('Formula: Ad = L x W + 2 x (3H) x (L + W) + pi x (3H)^2')
        self.doc.add_paragraph('Reference: IEC 62305-2 Annex A.2.1.1')
        p = self.doc.add_paragraph()
        p.add_run('Result: ').bold = True
        p.add_run(f'Ad = {results["ad"]:.2f} m²')
        
        self.doc.add_heading('1.2 Near strike collection area (Am)', level=1)
        self.doc.add_paragraph('Formula: Am = 2 x 500 x (L + W) + pi x 500^2')
        self.doc.add_paragraph('Reference: IEC 62305-2 Annex A.3')
        p = self.doc.add_paragraph()
        p.add_run('Result: ').bold = True
        p.add_run(f'Am = {results["am"]:.2f} m²')
        
        self.doc.add_heading('1.3 Environmental factor (CD)', level=1)
        self.doc.add_paragraph(f'Selected environment: {inputs.get("environment", "Isolated")}')
        p = self.doc.add_paragraph()
        p.add_run('Result: ').bold = True
        p.add_run(f'CD = {inputs.get("cd", 1)}')
        
        self.doc.add_heading('1.4 Lightning ground flash density (NG)', level=1)
        self.doc.add_paragraph('Formula: NG = 0.1 x Td')
        p = self.doc.add_paragraph()
        p.add_run('Result: ').bold = True
        p.add_run(f'NG = {results.get("ng", 1)} flashes/km²/year')
        
        self.doc.add_heading('1.5 Lightning frequencies', level=1)
        p = self.doc.add_paragraph()
        p.add_run('Nd (Direct): ').bold = True
        p.add_run(f'{results.get("nd", 0):.6f} events/year')
        p = self.doc.add_paragraph()
        p.add_run('Nm (Near): ').bold = True
        p.add_run(f'{results.get("nm", 0):.6f} events/year')
        
        self.doc.add_heading('1.6 Protection level', level=1)
        p = self.doc.add_paragraph()
        p.add_run('Result: ').bold = True
        p.add_run(f'{results.get("lpl", "Class III")}')
        self.doc.add_paragraph(f'Rolling sphere radius: {results.get("sphere", 45)}m')
        
        self.doc.add_page_break()
        self.doc.add_heading('SUMMARY OF RESULTS', level=1)
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'
        
        summary_data = [
            ('Collection area (Ad)', f"{results['ad']:.2f} m²"),
            ('Near strike area (Am)', f"{results['am']:.2f} m²"),
            ('Environmental factor (CD)', str(inputs.get('cd', 1))),
            ('Lightning density (NG)', f"{results.get('ng', 1)} flashes/km²/year"),
            ('Direct frequency (Nd)', f"{results.get('nd', 0):.6f} events/year"),
            ('Near frequency (Nm)', f"{results.get('nm', 0):.6f} events/year"),
            ('Protection efficiency', f"{results.get('efficiency', 0):.1%}"),
            ('Protection level', results.get('lpl', 'Class III')),
            ('Rolling sphere radius', f"{results.get('sphere', 45)} m"),
            ('Air terminals required', str(results.get('air_terminals', 4)))
        ]
        
        for param, value in summary_data:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = value
        
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(f'Generated by CES-Electrical on {format_pakistan_datetime()} (Pakistan Time)').italic = True
    
    def save(self, filename):
        self.doc.save(filename)

class CableWordReport:
    def __init__(self):
        self.doc = Document()
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        for section in self.doc.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
    
    def add_title(self):
        title = self.doc.add_heading('CABLE SIZING & CIRCUIT BREAKER SELECTION REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        p = self.doc.add_paragraph()
        p.add_run(f'Date: {format_pakistan_datetime()} (Pakistan Time)').italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.doc.add_paragraph()
    
    def add_common_parameters(self, ambient_temp):
        heading = self.doc.add_heading('COMMON PARAMETERS FOR ALL LOADS', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        params = [
            ('Ambient Temperature', f'{ambient_temp}°C'),
            ('Voltage Drop Limit', '2.5%'),
            ('Circuit Breaker Safety Factor', '1.25 (25%)'),
            ('Short Circuit Duration', '1 second'),
            ('Conductor Material', 'Copper'),
            ('Reference Standards', 'IEC 60364 / BS 7671 / IEC 60502-2')
        ]
        for i, (param, value) in enumerate(params):
            row_cells = table.rows[i]
            row_cells.cells[0].text = param
            row_cells.cells[1].text = value
            row_cells.cells[0].paragraphs[0].runs[0].bold = True
        self.doc.add_paragraph()
    
    def _get_load_value(self, load, param_name):
        try:
            if param_name == 'Power (kW)':
                return f"{float(load['Power (kW)']):.1f}"
            elif param_name == 'Voltage (V)':
                return f"{float(load['Voltage (V)']):.0f}"
            elif param_name == 'Phase':
                return str(load['Phase'])
            elif param_name == 'Load Type':
                return format_load_type(load.get('Load Type', 'Continuous'))
            elif param_name == 'Power Factor':
                return f"{float(load.get('Power Factor', 0.85)):.2f}"
            elif param_name == 'Efficiency':
                return f"{float(load.get('Efficiency', 1.0)):.2f}"
            elif param_name == 'Cable Length (m)':
                return f"{float(load['Length (m)']):.0f}"
            elif param_name == 'Insulation Type':
                return format_insulation_type(load.get('Insulation Type', 'XLPE_90'))
            elif param_name == 'Cable Type':
                return format_cable_type(load.get('Cable Type', 'multi_core_non_armoured'))
            elif param_name == 'Installation Method':
                return format_installation_method(load.get('Installation Method', 'C'))
            elif param_name == 'Cable Configuration':
                config_key = load.get('Table_Config', 'N/A')
                return get_table_config_description(config_key, load.get('Cable Type', ''))
            elif param_name == 'No. of Circuits':
                return str(load.get('Cables in Group', 1))
            elif param_name == 'Cable Formation':
                return format_cable_formation(load.get('Cable Formation', 'flat'))
            else:
                return 'N/A'
        except:
            return 'N/A'
    
    def add_load_details(self, loads_df):
        heading = self.doc.add_heading('LOAD DETAILS', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        if loads_df.empty:
            self.doc.add_paragraph('No load data available.')
            return
        
        CHUNK_SIZE = 15
        num_loads = len(loads_df)
        num_chunks = (num_loads + CHUNK_SIZE - 1) // CHUNK_SIZE
        
        param_names = [
            'Power (kW)', 'Voltage (V)', 'Phase', 'Load Type', 'Power Factor', 
            'Efficiency', 'Cable Length (m)', 'Insulation Type', 'Cable Type', 
            'Installation Method', 'Cable Configuration', 'No. of Circuits', 'Cable Formation'
        ]
        
        for chunk_idx in range(num_chunks):
            start_idx = chunk_idx * CHUNK_SIZE
            end_idx = min(start_idx + CHUNK_SIZE, num_loads)
            chunk_df = loads_df.iloc[start_idx:end_idx]
            
            if num_chunks > 1:
                self.doc.add_heading(f'Loads {start_idx + 1} to {end_idx}', level=2)
            
            num_cols = len(chunk_df) + 1
            num_rows = len(param_names) + 1
            table = self.doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            header_row = table.rows[0]
            header_row.cells[0].text = 'Parameter'
            header_row.cells[0].paragraphs[0].runs[0].bold = True
            
            for col_idx, (idx, load) in enumerate(chunk_df.iterrows(), start=1):
                header_row.cells[col_idx].text = load['Load Name']
                header_row.cells[col_idx].paragraphs[0].runs[0].bold = True
            
            for row_idx, param_name in enumerate(param_names, start=1):
                param_cell = table.rows[row_idx].cells[0]
                param_cell.text = param_name
                param_cell.paragraphs[0].runs[0].bold = True
                
                for col_idx, (idx, load) in enumerate(chunk_df.iterrows(), start=1):
                    value = self._get_load_value(load, param_name)
                    table.rows[row_idx].cells[col_idx].text = value
            
            self.doc.add_paragraph()
            
            burial_params = ['Soil Resistivity (K.m/W)', 'Burial Depth (m)']
            has_burial = any(load.get('Installation Method') in ['D', 'D_direct'] for _, load in chunk_df.iterrows())
            
            if has_burial:
                burial_table = self.doc.add_table(rows=len(chunk_df) + 1, cols=len(burial_params) + 1)
                burial_table.style = 'Light Grid Accent 1'
                
                burial_header = burial_table.rows[0]
                burial_header.cells[0].text = 'Load Name'
                burial_header.cells[0].paragraphs[0].runs[0].bold = True
                
                for col_idx, param_name in enumerate(burial_params, start=1):
                    burial_header.cells[col_idx].text = param_name
                    burial_header.cells[col_idx].paragraphs[0].runs[0].bold = True
                
                for row_idx, (idx, load) in enumerate(chunk_df.iterrows(), start=1):
                    row_cells = burial_table.rows[row_idx].cells
                    row_cells[0].text = load['Load Name']
                    
                    for col_idx, param_name in enumerate(burial_params, start=1):
                        if load.get('Installation Method') in ['D', 'D_direct']:
                            if param_name == 'Soil Resistivity (K.m/W)':
                                value = f"{load.get('Soil Resistivity (K.m/W)', 1.5)}"
                            else:
                                value = f"{load.get('Burial Depth (m)', 0.8)} m"
                        else:
                            value = 'N/A'
                        row_cells[col_idx].text = value
                
                self.doc.add_paragraph()
            
            if chunk_idx < num_chunks - 1:
                self.doc.add_page_break()
    
    def add_cable_results(self, cable_df):
        heading = self.doc.add_heading('CABLE SIZING RESULTS', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        if cable_df.empty:
            self.doc.add_paragraph('No cable results available.')
            return
        
        CHUNK_SIZE = 20
        num_loads = len(cable_df)
        num_chunks = (num_loads + CHUNK_SIZE - 1) // CHUNK_SIZE
        
        result_params = [
            'Size (mm²)', 'Load Current (A)', 'Current Carrying Capacity (A)',
            'Derating Factor K', 'Derated Ampacity (A)', 'Voltage Drop (%)',
            'Short Circuit (kA)', 'Status', 'Check'
        ]
        
        for chunk_idx in range(num_chunks):
            start_idx = chunk_idx * CHUNK_SIZE
            end_idx = min(start_idx + CHUNK_SIZE, num_loads)
            chunk_df = cable_df.iloc[start_idx:end_idx]
            
            if num_chunks > 1:
                self.doc.add_heading(f'Results for Loads {start_idx + 1} to {end_idx}', level=2)
            
            num_cols = len(result_params) + 1
            num_rows = len(chunk_df) + 1
            table = self.doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            header_row = table.rows[0]
            header_row.cells[0].text = 'Load Name'
            header_row.cells[0].paragraphs[0].runs[0].bold = True
            
            for col_idx, param in enumerate(result_params, start=1):
                header_row.cells[col_idx].text = param
                header_row.cells[col_idx].paragraphs[0].runs[0].bold = True
            
            for row_idx, (idx, cable_row) in enumerate(chunk_df.iterrows(), start=1):
                row_cells = table.rows[row_idx].cells
                row_cells[0].text = cable_row['Load Name']
                
                for col_idx, param in enumerate(result_params, start=1):
                    value = str(cable_row.get(param, 'N/A'))
                    cell = row_cells[col_idx]
                    cell.text = value
                    
                    if param == 'Status':
                        if value == 'PASS':
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                        else:
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    elif param == 'Check':
                        if value == 'PASS':
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                        else:
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            self.doc.add_paragraph()
            
            if chunk_idx < num_chunks - 1:
                self.doc.add_page_break()
    
    def add_detailed_calculations(self, detailed_calcs):
        self.doc.add_page_break()
        heading = self.doc.add_heading('DETAILED CABLE CALCULATIONS', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        for i, calc in enumerate(detailed_calcs):
            if i > 0:
                self.doc.add_paragraph()
            
            self.doc.add_heading(f'Load {i+1}: {calc["load_name"]} ({format_load_type(calc.get("load_type", "Continuous"))})', level=2)
            
            # Step 1: Load current calculation
            self.doc.add_heading('Step 1: Load current calculation', level=3)
            if calc['phase'] == '3-phase':
                p = self.doc.add_paragraph()
                p.add_run('Formula: I = P x 1000 / (1.732 x V x PF) for 3-phase').bold = True
                p = self.doc.add_paragraph()
                p.add_run('Calculation: ').bold = True
                p.add_run(f'I = {calc["power"]:.1f} x 1000 / (1.732 x {calc["voltage"]:.0f} x {calc["pf"]:.2f}) = {calc["current"]:.1f} A')
            else:
                p = self.doc.add_paragraph()
                p.add_run('Formula: I = P x 1000 / (V x PF) for 1-phase').bold = True
                p = self.doc.add_paragraph()
                p.add_run('Calculation: ').bold = True
                p.add_run(f'I = {calc["power"]:.1f} x 1000 / ({calc["voltage"]:.0f} x {calc["pf"]:.2f}) = {calc["current"]:.1f} A')
            
            # Step 2: Cable type and reference selection
            self.doc.add_heading('Step 2: Cable type and reference selection', level=3)
            p = self.doc.add_paragraph()
            p.add_run(f'Voltage {calc["voltage"]:.0f}V -> {calc["cable_category"]} cables selected')
            p = self.doc.add_paragraph()
            p.add_run(f'Cable Type: {format_cable_type(calc["cable_type"])}')
            p.add_run(f' | Formation: {format_cable_formation(calc["formation"])}')
            p = self.doc.add_paragraph()
            p.add_run(f'Reference Method: {format_installation_method(calc["installation"])}')
            
            # Step 3: Derating factors calculation
            self.doc.add_heading('Step 3: Derating factors calculation', level=3)
            self.doc.add_paragraph('Total derating factor K = k1 × k2 × k3 × k4')
            
            factor_table = self.doc.add_table(rows=5, cols=2)
            factor_table.style = 'Light Grid Accent 1'
            factors_data = [
                ('k1 (Temperature correction)', f'{calc["k1"]:.3f} - at {calc["ambient_temp"]}°C'),
                ('k2 (Grouping correction)', f'{calc["k2"]:.3f} - {format_cable_arrangement(calc["arrangement"])}'),
                ('k3 (Soil resistivity correction)', f'{calc["k3"]:.3f}'),
                ('k4 (Depth correction)', f'{calc["k4"]:.3f}'),
                ('Total K', f'{calc["total_k"]:.3f}')
            ]
            for j, (param, value) in enumerate(factors_data):
                row_cells = factor_table.rows[j]
                row_cells.cells[0].text = param
                row_cells.cells[1].text = value
                row_cells.cells[0].paragraphs[0].runs[0].bold = True
            
            # Step 4: Cable selection
            self.doc.add_heading('Step 4: Cable selection (automatic)', level=3)
            p = self.doc.add_paragraph()
            p.add_run('Selected cable: ').bold = True
            p.add_run(f'{calc["size"]} mm² {format_cable_type(calc["cable_type"])}')
            
            ampacity_table = self.doc.add_table(rows=3, cols=2)
            ampacity_table.style = 'Light Grid Accent 1'
            ampacity_data = [
                ('Current Carrying Capacity (from table)', f'{calc["base_amp"]} A'),
                ('Derated ampacity = CCC × K', f'{calc["base_amp"]:.0f} × {calc["total_k"]:.3f} = {calc["derated_amp"]:.1f} A'),
                ('Ampacity Check', f'{calc["derated_amp"]:.1f} A >= {calc["current"]:.1f} A ? {"PASS" if calc["ampacity_pass"] else "FAIL"}')
            ]
            for j, (param, value) in enumerate(ampacity_data):
                row_cells = ampacity_table.rows[j]
                row_cells.cells[0].text = param
                row_cells.cells[1].text = value
                row_cells.cells[0].paragraphs[0].runs[0].bold = True
                if param == 'Ampacity Check' and calc['ampacity_pass']:
                    row_cells.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                elif param == 'Ampacity Check':
                    row_cells.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            # Step 5: Voltage drop calculation
            self.doc.add_heading('Step 5: Voltage drop calculation', level=3)
            
            vd_values = get_voltage_drop_values(calc['cable_type'], calc['size'], calc['phase'], calc['formation'])
            
            self.doc.add_paragraph(f'Using values from BS 7671 Table for {format_cable_type(calc["cable_type"])}')
            
            if vd_values['type'] == 'mv':
                mv_value = vd_values['value']
                self.doc.add_paragraph(f'From table: mV/A/m = {mv_value}')
                self.doc.add_paragraph('Formula: Vd = (mV/A/m) × I × L / 1000')
                self.doc.add_paragraph(f'Vd = {mv_value} × {calc["current"]:.1f} × {calc["length"]:.0f} / 1000')
                vd_calc = mv_value * calc['current'] * calc['length'] / 1000
                self.doc.add_paragraph(f'Vd = {vd_calc:.2f} V')
                self.doc.add_paragraph(f'Vd% = ({vd_calc:.2f} / {calc["voltage"]:.0f}) × 100 = {calc["vd_pct"]:.3f}%')
            else:
                r = vd_values['R']
                x = vd_values['X']
                phi = math.acos(calc['pf'])
                sin_phi = math.sin(phi)
                
                self.doc.add_paragraph(f'From table: R = {r:.4f} Ω/km, X = {x:.4f} Ω/km')
                self.doc.add_paragraph(f'Power factor cosφ = {calc["pf"]:.3f}, sinφ = √(1 - cos²φ) = {sin_phi:.4f}')
                self.doc.add_paragraph(f'R cosφ = {r:.4f} × {calc["pf"]:.3f} = {r * calc["pf"]:.4f}')
                self.doc.add_paragraph(f'X sinφ = {x:.4f} × {sin_phi:.4f} = {x * sin_phi:.4f}')
                self.doc.add_paragraph(f'(R cosφ + X sinφ) = {r * calc["pf"] + x * sin_phi:.4f}')
                
                if calc['phase'] == '3-phase':
                    self.doc.add_paragraph('Formula (3-phase): Vd = √3 × I × (R cosφ + X sinφ) × L / 1000')
                    vd_calc = 1.732 * calc['current'] * (r * calc['pf'] + x * sin_phi) * calc['length'] / 1000
                    self.doc.add_paragraph(f'Vd = 1.732 × {calc["current"]:.1f} × {(r * calc["pf"] + x * sin_phi):.4f} × {calc["length"]:.0f} / 1000')
                    self.doc.add_paragraph(f'Vd = {vd_calc:.2f} V')
                else:
                    self.doc.add_paragraph('Formula (1-phase): Vd = 2 × I × (R cosφ + X sinφ) × L / 1000')
                    vd_calc = 2 * calc['current'] * (r * calc['pf'] + x * sin_phi) * calc['length'] / 1000
                    self.doc.add_paragraph(f'Vd = 2 × {calc["current"]:.1f} × {(r * calc["pf"] + x * sin_phi):.4f} × {calc["length"]:.0f} / 1000')
                    self.doc.add_paragraph(f'Vd = {vd_calc:.2f} V')
                
                self.doc.add_paragraph(f'Vd% = ({vd_calc:.2f} / {calc["voltage"]:.0f}) × 100 = {calc["vd_pct"]:.3f}%')
            
            vd_table = self.doc.add_table(rows=3, cols=2)
            vd_table.style = 'Light Grid Accent 1'
            vd_data = [
                ('Calculated Voltage Drop', f'{calc["vd_pct"]:.3f}%'),
                ('Limit', '2.5%'),
                ('VD Check', f'{calc["vd_pct"]:.3f}% <= 2.5% ? {"PASS" if calc["vd_pass"] else "FAIL"}')
            ]
            for j, (param, value) in enumerate(vd_data):
                row_cells = vd_table.rows[j]
                row_cells.cells[0].text = param
                row_cells.cells[1].text = value
                row_cells.cells[0].paragraphs[0].runs[0].bold = True
                if param == 'VD Check' and calc['vd_pass']:
                    row_cells.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                elif param == 'VD Check':
                    row_cells.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            # Step 6: Short circuit calculation
            self.doc.add_heading('Step 6: Short circuit calculation', level=3)
            self.doc.add_paragraph('Formula: Isc = (K × S / √t) × √(ln((θf + β) / (θi + β)))')
            self.doc.add_paragraph(f'Where: K = 226 (Copper), β = 234.5, θi = {calc["theta_i"]:.0f}°C, θf = {calc["theta_f"]:.0f}°C')
            self.doc.add_paragraph(f'First term = 226 × {calc["size"]} / √1.0 = {226 * calc["size"]:.1f}')
            self.doc.add_paragraph(f'ln term = ln(({calc["theta_f"]:.0f} + 234.5) / ({calc["theta_i"]:.0f} + 234.5)) = {math.log((calc["theta_f"] + 234.5) / (calc["theta_i"] + 234.5)):.4f}')
            self.doc.add_paragraph(f'√(ln term) = {math.sqrt(math.log((calc["theta_f"] + 234.5) / (calc["theta_i"] + 234.5))):.4f}')
            self.doc.add_paragraph(f'Isc = {calc["sc"]:.2f} kA')
            
            # Final status
            self.doc.add_heading('Final status', level=3)
            p = self.doc.add_paragraph()
            if calc['status'] == 'PASS':
                final = p.add_run('PASS')
                final.font.color.rgb = RGBColor(0, 128, 0)
            else:
                final = p.add_run('FAIL')
                final.font.color.rgb = RGBColor(255, 0, 0)
            final.font.size = Pt(14)
            final.font.bold = True
            self.doc.add_paragraph('_' * 60)
    
    def add_cb_results(self, cb_results, main_cbs_by_voltage, pole_selections, main_pole_selections, cb_details=None, selected_manufacturer='Schneider Electric'):
        self.doc.add_page_break()
        heading = self.doc.add_heading('CIRCUIT BREAKER SIZING', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        if cb_details:
            self.doc.add_heading('DETAILED CIRCUIT BREAKER CALCULATIONS', level=2)
            for i, detail in enumerate(cb_details):
                self.doc.add_heading(f'Load {i+1}: {detail["load_name"]}', level=3)
                
                self.doc.add_heading('Step 1: Load current calculation', level=4)
                p = self.doc.add_paragraph()
                if detail['phase_desc'] == "Three-phase":
                    p.add_run('Formula: I = P x 1000 / (1.732 x V x PF)').bold = True
                    p = self.doc.add_paragraph()
                    p.add_run('Calculation: ').bold = True
                    orig_power = detail['current'] * 1.732 * detail['voltage'] * 0.85 / 1000
                    p.add_run(f'I = {orig_power:.1f} x 1000 / (1.732 x {detail["voltage"]:.0f} x 0.85) = {detail["current"]:.2f} A')
                elif detail['phase_desc'] == "Single-phase":
                    p.add_run('Formula: I = P x 1000 / (V x PF)').bold = True
                    p = self.doc.add_paragraph()
                    p.add_run('Calculation: ').bold = True
                    orig_power = detail['current'] * detail['voltage'] * 0.85 / 1000
                    p.add_run(f'I = {orig_power:.1f} x 1000 / ({detail["voltage"]:.0f} x 0.85) = {detail["current"]:.2f} A')
                else:
                    p.add_run('Formula: I = P x 1000 / V (DC)').bold = True
                    p = self.doc.add_paragraph()
                    p.add_run('Calculation: ').bold = True
                    orig_power = detail['current'] * 110 / 1000
                    p.add_run(f'I = {orig_power:.1f} x 1000 / 110 = {detail["current"]:.2f} A')
                
                self.doc.add_heading('Step 2: Circuit breaker rating calculation', level=4)
                self.doc.add_paragraph(f'Safety factor: {detail["design_factor"]} (25% safety margin)')
                p = self.doc.add_paragraph()
                p.add_run('Required rating = Required current × Safety factor').bold = True
                p = self.doc.add_paragraph()
                p.add_run(f'Required = {detail["current"]:.2f} × {detail["design_factor"]} = {detail["required"]:.2f} A')
                
                self.doc.add_heading('Step 3: Final selection', level=4)
                p = self.doc.add_paragraph()
                p.add_run(f'Selected Circuit Breaker Rating: {detail["selected"]} A').bold = True
                self.doc.add_paragraph('_' * 50)
        
        self.doc.add_heading('INDIVIDUAL CIRCUIT BREAKERS SUMMARY', level=2)
        
        cb_params = ['Power (kW)', 'Voltage (V)', 'Phase', 'Load Current (A)', 
                     'Selected CB (A)', 'Breaker Type', 'Poles']
        
        CHUNK_SIZE = 20
        num_loads = len(cb_results)
        num_chunks = (num_loads + CHUNK_SIZE - 1) // CHUNK_SIZE
        
        for chunk_idx in range(num_chunks):
            start_idx = chunk_idx * CHUNK_SIZE
            end_idx = min(start_idx + CHUNK_SIZE, num_loads)
            chunk_results = cb_results[start_idx:end_idx]
            
            if num_chunks > 1:
                self.doc.add_heading(f'Breakers for Loads {start_idx + 1} to {end_idx}', level=3)
            
            num_cols = len(cb_params) + 1
            num_rows = len(chunk_results) + 1
            table = self.doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            header_row = table.rows[0]
            header_row.cells[0].text = 'Load Name'
            header_row.cells[0].paragraphs[0].runs[0].bold = True
            
            for col_idx, param in enumerate(cb_params, start=1):
                header_row.cells[col_idx].text = param
                header_row.cells[col_idx].paragraphs[0].runs[0].bold = True
            
            for row_idx, r in enumerate(chunk_results, start=1):
                row_cells = table.rows[row_idx].cells
                row_cells[0].text = r['Load']
                
                for col_idx, param in enumerate(cb_params, start=1):
                    if param == 'Power (kW)':
                        value = f"{r['Power (kW)']:.1f}"
                    elif param == 'Voltage (V)':
                        value = f"{r['Voltage (V)']:.0f}"
                    elif param == 'Phase':
                        value = r['Phase']
                    elif param == 'Load Current (A)':
                        value = f"{r['Current (A)']:.1f}"
                    elif param == 'Selected CB (A)':
                        value = str(r['Selected CB (A)'])
                    elif param == 'Breaker Type':
                        value = f"{r['Breaker Type']}"
                    elif param == 'Poles':
                        value = pole_selections.get(r['Load'], '3P')
                    else:
                        value = 'N/A'
                    row_cells[col_idx].text = value
            
            self.doc.add_paragraph()
            
            if chunk_idx < num_chunks - 1:
                self.doc.add_page_break()
        
        self.doc.add_heading('MAIN CIRCUIT BREAKERS (BY VOLTAGE LEVEL)', level=2)
        
        if main_cbs_by_voltage:
            for voltage, main_cb in main_cbs_by_voltage.items():
                self.doc.add_heading(f'{main_cb["system_type"]} - {main_cb["voltage_range"]} System', level=3)
                
                main_table = self.doc.add_table(rows=9, cols=2)
                main_table.style = 'Light Grid Accent 1'
                selected_poles = main_pole_selections.get(f'voltage_{voltage}', '3P')
                main_params = [
                    ('System Voltage', f"{main_cb['voltage']:.0f} V ({main_cb['voltage_range']})"),
                    ('Number of Loads', str(main_cb['num_loads'])),
                    ('Total Power', f"{main_cb['total_power']:.1f} kW"),
                    ('Average Power Factor', f"{main_cb['avg_pf']:.3f}"),
                    ('Total Current', f"{main_cb['current']:.1f} A"),
                    ('Required CB Rating', f"{main_cb['required_cb']:.1f} A"),
                    ('Selected CB Rating', f"{main_cb['selected_cb']} A"),
                    ('Breaker Type', f"{main_cb['breaker_type']} ({main_cb['standard']})"),
                    ('Poles', selected_poles)
                ]
                for i, (param, value) in enumerate(main_params):
                    if i < len(main_table.rows):
                        row_cells = main_table.rows[i]
                        row_cells.cells[0].text = param
                        row_cells.cells[1].text = value
                        row_cells.cells[0].paragraphs[0].runs[0].bold = True
                
                self.doc.add_heading(f'Detailed Calculation for {main_cb["voltage_range"]} Main Circuit Breaker', level=4)
                
                loads_list = ""
                for r in cb_results:
                    if r.get('Voltage (V)', 0) == voltage:
                        pf_val = r.get('Power Factor', 0.85)
                        loads_list += f"  - {r.get('Load', 'Unknown')}: {r.get('Power (kW)', 0):.1f} kW, PF={pf_val}\n"
                
                self.doc.add_paragraph(f'Step 1: Load analysis for {main_cb["voltage_range"]} system')
                self.doc.add_paragraph(f'Voltage level: {main_cb["voltage"]:.0f} V ({main_cb["system_type"]})')
                self.doc.add_paragraph(f'Number of loads in this group: {main_cb["num_loads"]}')
                self.doc.add_paragraph('Loads in this voltage group:')
                for line in loads_list.split('\n'):
                    if line.strip():
                        self.doc.add_paragraph(line, style='List Bullet')
                
                self.doc.add_paragraph(f'Total connected load: {main_cb["total_power"]:.2f} kW')
                self.doc.add_paragraph(f'Weighted average power factor: {main_cb["avg_pf"]:.3f}')
                
                self.doc.add_paragraph(f'Step 2: Total current calculation')
                self.doc.add_paragraph(f'Formula: I = P x 1000 / (1.732 x V x PF)')
                self.doc.add_paragraph(f'I = {main_cb["total_power"]:.2f} x 1000 / (1.732 x {main_cb["voltage"]:.0f} x {main_cb["avg_pf"]:.3f})')
                self.doc.add_paragraph(f'I = {main_cb["current"]:.2f} A')
                
                self.doc.add_paragraph(f'Step 3: Circuit breaker sizing')
                self.doc.add_paragraph(f'Safety factor: 1.25 (25% safety margin for continuous loads)')
                self.doc.add_paragraph(f'Required rating = Required current × Safety factor')
                self.doc.add_paragraph(f'Required = {main_cb["current"]:.2f} × 1.25 = {main_cb["required_cb"]:.2f} A')
                
                self.doc.add_paragraph(f'Step 4: Final selection')
                p = self.doc.add_paragraph()
                p.add_run(f'Selected Circuit Breaker Rating: {main_cb["selected_cb"]} A').bold = True
                
                self.doc.add_paragraph('_' * 60)
        else:
            self.doc.add_paragraph('No main circuit breaker calculations available.')
    
    def save(self, filename):
        self.doc.save(filename)

class TransformerWordReport:
    def __init__(self):
        self.doc = Document()
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
    
    def add_title(self):
        title = self.doc.add_heading('TRANSFORMER SIZING REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        self.doc.add_paragraph(f'Date: {format_pakistan_datetime()} (Pakistan Time)')
        self.doc.add_paragraph()
    
    def add_load_analysis(self, loads_df):
        heading = self.doc.add_heading('LOAD ANALYSIS', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        table = self.doc.add_table(rows=1, cols=8)
        table.style = 'Light Grid Accent 1'
        headers = ['Load description', 'Qty', 'Rating (kw)', 'Voltage (V)', 'PF', 'Connected (kw)', 'Diversity', 'P (kw)']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        total_p = 0
        for idx, load in loads_df.iterrows():
            load_type_diversity = LOAD_TYPE_FACTORS[load['Load Type']]['diversity']
            connected = load['Rating (kW)'] * load['Quantity']
            p = connected * load_type_diversity
            total_p += p
            row = table.add_row().cells
            row[0].text = load['Load Description']
            row[1].text = str(load['Quantity'])
            row[2].text = f"{load['Rating (kW)']:.0f}"
            row[3].text = f"{load['Voltage (V)']:.0f}"
            row[4].text = f"{load['Power Factor']:.2f}"
            row[5].text = f"{connected:.0f}"
            row[6].text = f"{load_type_diversity:.1f}"
            row[7].text = f"{p:.1f}"
        p_row = table.add_row().cells
        p_row[0].text = 'Total real power (P)'
        p_row[0].paragraphs[0].runs[0].bold = True
        for i in range(1, 7):
            p_row[i].text = ''
        p_row[7].text = f"{total_p:.1f} kW"
        p_row[7].paragraphs[0].runs[0].bold = True
        self.doc.add_paragraph()
        return total_p
    
    def add_step_by_step(self, loads_df):
        self.doc.add_page_break()
        heading = self.doc.add_heading('STEP-BY-STEP P, Q, S CALCULATIONS', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        total_p = 0
        total_q = 0
        for idx, load in loads_df.iterrows():
            self.doc.add_heading(f'Load {idx+1}: {load["Load Description"]}', level=2)
            load_type_diversity = LOAD_TYPE_FACTORS[load['Load Type']]['diversity']
            connected = load['Rating (kW)'] * load['Quantity']
            p = connected * load_type_diversity
            phi = math.acos(load['Power Factor'])
            tan_phi = math.tan(phi)
            q = p * tan_phi
            s = math.sqrt(p**2 + q**2)
            self.doc.add_paragraph(f'Step 1 - Connected power: {load["Rating (kW)"]:.0f} kW x {load["Quantity"]} = {connected:.0f} kW')
            self.doc.add_paragraph(f'Step 2 - Demand power (P): {connected:.0f} kW x {load_type_diversity:.1f} = {p:.1f} kW')
            self.doc.add_paragraph(f'Step 3 - Angle φ: acos({load["Power Factor"]}) = {math.degrees(phi):.1f}°')
            self.doc.add_paragraph(f'Step 4 - tan(φ): tan({math.degrees(phi):.1f}°) = {tan_phi:.3f}')
            self.doc.add_paragraph(f'Step 5 - Reactive power (Q): {p:.1f} kW x {tan_phi:.3f} = {q:.1f} kVAR')
            p_step = self.doc.add_paragraph()
            p_step.add_run('Step 6 - Apparent power (S): ').bold = True
            p_step.add_run(f'√({p:.1f}² + {q:.1f}²) = {s:.1f} kVA')
            total_p += p
            total_q += q
            self.doc.add_paragraph('_' * 50)
        return total_p, total_q
    
    def add_largest_equipment(self, loads_df, total_p, total_s, largest_data):
        self.doc.add_page_break()
        heading = self.doc.add_heading('LARGEST EQUIPMENT ANALYSIS', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        if largest_data:
            self.doc.add_heading(f'Largest equipment: {largest_data["load"]["Load Description"]}', level=2)
            self.doc.add_paragraph(f'Connected power: {largest_data["connected"]:.0f} kW')
            self.doc.add_paragraph(f'Demand power (P): {largest_data["p"]:.1f} kW')
            self.doc.add_paragraph(f'Reactive power (Q): {largest_data["q"]:.1f} kVAR')
            self.doc.add_paragraph(f'Apparent power (S): {largest_data["s"]:.1f} kVA')
            p_pct = (largest_data["p"] / total_p) * 100 if total_p > 0 else 0
            s_pct = (largest_data["s"] / total_s) * 100 if total_s > 0 else 0
            self.doc.add_heading('Impact on total system:', level=3)
            self.doc.add_paragraph(f'• Contributes {p_pct:.1f}% of total real power (P)')
            self.doc.add_paragraph(f'• Contributes {s_pct:.1f}% of total apparent power (S)')
    
    def add_transformer_selection(self, total_p, total_q, future_expansion, selected_kva):
        self.doc.add_page_break()
        heading = self.doc.add_heading('TRANSFORMER SELECTION', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        total_s = math.sqrt(total_p**2 + total_q**2)
        with_future = total_s * (1 + future_expansion/100)
        self.doc.add_paragraph(f'Total real power (P) = {total_p:.1f} kW')
        self.doc.add_paragraph(f'Total reactive power (Q) = {total_q:.1f} kVAR')
        p = self.doc.add_paragraph()
        p.add_run(f'Total apparent power (S) = √({total_p:.1f}² + {total_q:.1f}²) = {total_s:.1f} kVA').bold = True
        self.doc.add_paragraph()
        self.doc.add_paragraph(f'Future expansion: +{future_expansion}%')
        self.doc.add_paragraph(f'Required with future = {total_s:.1f} x {1 + future_expansion/100:.2f} = {with_future:.1f} kVA')
        self.doc.add_paragraph()
        self.doc.add_heading('Standard ratings:', level=3)
        ratings = [50, 100, 160, 200, 250, 315, 400, 500, 630, 800, 1000, 1250, 1600, 2000, 2500, 3150]
        self.doc.add_paragraph(', '.join(str(r) for r in ratings))
        final_heading = self.doc.add_heading('', level=2)
        final_heading.add_run(f'Selected transformer: {selected_kva} kVA').bold = True
        final_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def save(self, filename):
        self.doc.save(filename)

class SimpleTransformerCalculator:
    def __init__(self):
        self.standard_ratings = [50, 100, 160, 200, 250, 315, 400, 500, 630, 800, 
                                  1000, 1250, 1600, 2000, 2500, 3150, 4000, 5000, 
                                  6300, 8000, 10000, 12500, 16000, 20000, 25000, 
                                  31500, 40000, 50000, 63000]
    
    def calculate_p(self, rating_kw, quantity, diversity):
        return rating_kw * quantity * diversity
    
    def calculate_q(self, p_kw, pf):
        if pf >= 1.0:
            return 0
        phi = math.acos(pf)
        return p_kw * math.tan(phi)
    
    def calculate_s(self, p_kw, q_kvar):
        return math.sqrt(p_kw**2 + q_kvar**2)
    
    def get_standard_rating(self, required_kva):
        for rating in self.standard_ratings:
            if rating >= required_kva:
                return rating
        return self.standard_ratings[-1]
    
    def find_largest_equipment(self, loads):
        max_p = 0
        max_load = None
        max_idx = -1
        for idx, load in loads.iterrows():
            p_connected = load['Rating (kW)'] * load['Quantity']
            if p_connected > max_p:
                max_p = p_connected
                max_load = load
                max_idx = idx
        return max_idx, max_load, max_p

# ========== SESSION STATE INITIALIZATION ==========
if 'universal_loads' not in st.session_state:
    st.session_state.universal_loads = pd.DataFrame({
        'Load Description': ['LV Motor 1', 'MV Motor 1'],
        'Quantity': [1, 1],
        'Rating (kW)': [75, 500],
        'Voltage (V)': [415, 3300],
        'Power Factor': [0.85, 0.85],
        'Load Type': ['Continuous', 'Continuous'],
        'Diversity Factor': [0.8, 0.8]
    })

if 'loads_df' not in st.session_state:
    st.session_state.loads_df = pd.DataFrame(columns=[
        'Load Name', 'Power (kW)', 'Voltage (V)', 'Phase', 'Load Type', 
        'Power Factor', 'Efficiency', 'Length (m)', 'Insulation Type', 
        'Cable Type', 'Installation Method', 'Table_Config', 'Cables in Group', 
        'Cable Arrangement', 'Cable Formation', 'Cable Clearance', 'Soil Resistivity (K.m/W)', 
        'Burial Depth (m)'
    ])

if 'cable_results_df' not in st.session_state:
    st.session_state.cable_results_df = pd.DataFrame()
if 'detailed_calcs' not in st.session_state:
    st.session_state.detailed_calcs = []
if 'all_derating_factors' not in st.session_state:
    st.session_state.all_derating_factors = {}
if 'cb_results' not in st.session_state:
    st.session_state.cb_results = []
if 'cb_details' not in st.session_state:
    st.session_state.cb_details = []
if 'main_cbs_by_voltage' not in st.session_state:
    st.session_state.main_cbs_by_voltage = {}
if 'main_cb_details_by_voltage' not in st.session_state:
    st.session_state.main_cb_details_by_voltage = {}
if 'calc_results' not in st.session_state:
    st.session_state.calc_results = {}
if 'calc_done' not in st.session_state:
    st.session_state.calc_done = False
if 'total_p' not in st.session_state:
    st.session_state.total_p = 0
if 'total_q' not in st.session_state:
    st.session_state.total_q = 0
if 'tx_largest_data' not in st.session_state:
    st.session_state.tx_largest_data = None

# ========== SIDEBAR NAVIGATION ==========
with st.sidebar:
    st.markdown('<div style="background: linear-gradient(135deg, #1E3A8A 0%, #3B5BA6 100%); color: white; padding: 15px; border-radius: 10px; margin-bottom: 20px; text-align: center;"><h2 style="color: white !important; margin: 0;">CES-Electrical</h2></div>', unsafe_allow_html=True)
    if 'selected_calculator' not in st.session_state:
        st.session_state.selected_calculator = "LOAD SHEET"
    
    calculators = [
        "LOAD SHEET",
        "Lightning Protection",
        "Cable Sizing",
        "Transformer Sizing",
        "Generator Sizing",
        "Earthing"
    ]
    
    for calc in calculators:
        if st.button(calc, key=f"nav_{calc}", use_container_width=True):
            st.session_state.selected_calculator = calc
            st.rerun()

st.title(f"{st.session_state.selected_calculator} Calculator")

# ========== LOAD SHEET TAB ==========
if st.session_state.selected_calculator == "LOAD SHEET":
    st.markdown('<div class="report-header">📋 UNIVERSAL LOAD SHEET</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class="info-box">
        <h4>📌 This load sheet is used by all calculators</h4>
        <p>Two example loads provided: LV Motor 1 (415V) and MV Motor 1 (3300V). Load Type determines diversity factor:</p>
        <p><span class="load-type-badge continuous-badge">Continuous</span> = 100% (1.0) | 
           <span class="load-type-badge intermittent-badge">Intermittent</span> = 30% (0.3) | 
           <span class="load-type-badge standby-badge">Standby</span> = 10% (0.1)</p>
    </div>
    """, unsafe_allow_html=True)
    
    with st.expander("📊 tan(acos(PF)) reference table", expanded=False):
        tan_data = {
            'Power factor': [1.0, 0.95, 0.90, 0.85, 0.80, 0.75, 0.70],
            'tan(acos(PF))': [0.00, 0.33, 0.48, 0.62, 0.75, 0.88, 1.02],
            'Example': ['PF=1.0 -> Q=0', 'PF=0.95 -> Q=0.33xP', 'PF=0.90 -> Q=0.48xP', 
                       'PF=0.85 -> Q=0.62xP', 'PF=0.80 -> Q=0.75xP', 'PF=0.75 -> Q=0.88xP', 
                       'PF=0.70 -> Q=1.02xP']
        }
        tan_df = pd.DataFrame(tan_data)
        st.dataframe(tan_df, use_container_width=True, hide_index=True)
        st.caption("Formula: Q = P x tan(acos(PF)) as per document")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("➕ Add load", key="add_load_main", use_container_width=True):
            new_row = pd.DataFrame({
                'Load Description': [f'Load {len(st.session_state.universal_loads) + 1}'],
                'Quantity': [1],
                'Rating (kW)': [50.0],
                'Voltage (V)': [415],
                'Power Factor': [0.85],
                'Load Type': ['Continuous'],
                'Diversity Factor': [0.8]
            })
            st.session_state.universal_loads = pd.concat([st.session_state.universal_loads, new_row], ignore_index=True)
            st.rerun()
    with col2:
        if st.button("🗑️ Delete last load", key="delete_load_main", use_container_width=True):
            if len(st.session_state.universal_loads) > 1:
                st.session_state.universal_loads = st.session_state.universal_loads[:-1]
                st.rerun()
            else:
                st.warning("At least one row required")
    
    st.markdown("### 📋 Current loads")
    for idx, load in st.session_state.universal_loads.iterrows():
        load_type_factor = LOAD_TYPE_FACTORS[load['Load Type']]['diversity']
        badge_color = "continuous-badge" if load['Load Type'] == 'Continuous' else "intermittent-badge" if load['Load Type'] == 'Intermittent' else "standby-badge"
        cable_type = "LV" if load['Voltage (V)'] <= 1000 else "MV"
        st.markdown(f"""
        <div class="calc-step">
            <h4>📌 Load {idx+1}: <span style="color: #1E3A8A;">{load['Load Description']}</span></h4>
            <p><b>Quantity:</b> {load['Quantity']} | <b>Rating:</b> {load['Rating (kW)']} kW | <b>Voltage:</b> {load['Voltage (V)']}V ({cable_type})</p>
            <p><b>PF:</b> {load['Power Factor']} | <b>Load type:</b> <span class="load-type-badge {badge_color}">{load['Load Type']} ({load_type_factor*100:.0f}%)</span> | <b>Diversity:</b> {load['Diversity Factor']}</p>
        </div>
        """, unsafe_allow_html=True)
    
    edited_loads = st.data_editor(
        st.session_state.universal_loads,
        num_rows="fixed",
        use_container_width=True,
        column_config={
            "Load Description": st.column_config.TextColumn("Load description", width="medium"),
            "Quantity": st.column_config.NumberColumn("Qty", min_value=1, max_value=100, step=1),
            "Rating (kW)": st.column_config.NumberColumn("Rating (kw)", min_value=0.0, max_value=10000.0, step=1.0),
            "Voltage (V)": st.column_config.NumberColumn("Voltage (v)", min_value=0, max_value=11000, step=100),
            "Power Factor": st.column_config.NumberColumn("Pf", min_value=0.5, max_value=1.0, step=0.05),
            "Load Type": st.column_config.SelectboxColumn("Load type", options=['Continuous', 'Intermittent', 'Standby']),
            "Diversity Factor": st.column_config.NumberColumn("Diversity", min_value=0.0, max_value=1.0, step=0.05)
        }
    )
    st.session_state.universal_loads = edited_loads
    
    total_connected = 0
    total_p = 0
    lv_count = 0
    mv_count = 0
    summary_data = []
    for idx, load in st.session_state.universal_loads.iterrows():
        load_type_diversity = LOAD_TYPE_FACTORS[load['Load Type']]['diversity']
        connected = load['Rating (kW)'] * load['Quantity']
        p = connected * load_type_diversity
        total_connected += connected
        total_p += p
        if load['Voltage (V)'] <= 1000:
            lv_count += 1
        else:
            mv_count += 1
        summary_data.append({
            'Load': load['Load Description'],
            'Connected (kw)': f"{connected:.0f}",
            'Demand (kw)': f"{p:.0f} ({load['Load Type']} {load_type_diversity*100:.0f}%)",
            'Pf': f"{load['Power Factor']:.2f}",
            'Type': load['Load Type']
        })
    summary_df = pd.DataFrame(summary_data)
    st.dataframe(summary_df, use_container_width=True, hide_index=True)
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total connected", f"{total_connected:.0f} kW")
    with col2:
        st.metric("Total demand", f"{total_p:.0f} kW")
    with col3:
        st.metric("LV loads", lv_count)
    with col4:
        st.metric("MV loads", mv_count)

# ========== LIGHTNING PROTECTION TAB ==========
elif st.session_state.selected_calculator == "Lightning Protection":
    st.markdown('<div class="report-header">⚡ LIGHTNING PROTECTION CALCULATOR</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        length = st.number_input("Length (m)", value=26.5, step=0.5)
        width = st.number_input("Width (m)", value=26.25, step=0.5)
        height = st.number_input("Height (m)", value=7.35, step=0.5)
        td_days = st.number_input("Thunderstorm days/year", value=10, step=1)
    with col2:
        environment = st.selectbox("Environment", ["Surrounded", "Similar height", "Isolated", "Hilltop"])
        cd_values = {"Surrounded": 0.25, "Similar height": 0.5, "Isolated": 1, "Hilltop": 2}
        cd = cd_values[environment]
        st.metric("Environmental factor (CD)", cd)
    
    if st.button("🔧 Calculate risk", type="primary", use_container_width=True):
        ad = length * width + 2 * (3 * height) * (length + width) + math.pi * (3 * height)**2
        am = 2 * 500 * (length + width) + math.pi * 500**2
        ng = 0.1 * td_days
        nd = ng * ad * cd * 1e-6
        nm = ng * am * 1e-6
        c2, c3, c4, c5 = 1.0, 3.0, 1.0, 5.0
        c_total = cd * c2 * c3 * c4 * c5
        nc = 1e-4 / c_total
        efficiency = 1 - (nc / nd) if nd > 0 else 0
        if efficiency > 0.98:
            lpl, sphere = "Class I", 20
        elif efficiency > 0.95:
            lpl, sphere = "Class II", 30
        elif efficiency > 0.90:
            lpl, sphere = "Class III", 45
        else:
            lpl, sphere = "Class IV", 60
        
        st.markdown("---")
        col_a, col_b, col_c, col_d = st.columns(4)
        with col_a:
            st.metric("Collection area (Ad)", f"{ad:.0f} m²")
        with col_b:
            st.metric("Nd (Direct)", f"{nd:.6f}")
        with col_c:
            st.metric("Protection level", lpl)
        with col_d:
            st.metric("Rolling sphere", f"{sphere}m")
        
        st.session_state.calc_results = {
            'ad': ad, 'am': am, 'ng': ng, 'nd': nd, 'nm': nm,
            'efficiency': efficiency, 'lpl': lpl, 'sphere': sphere, 'air_terminals': 4
        }
        st.session_state.input_values = {
            'length': length, 'width': width, 'height': height,
            'td_days': td_days, 'environment': environment, 'cd': cd
        }
        st.session_state.calc_done = True
    
    if st.session_state.calc_done:
        if st.button("📥 Generate word report", key="lp_word", use_container_width=True):
            with st.spinner("Generating word report..."):
                try:
                    word = LightningWordReport()
                    word.add_calculations(st.session_state.calc_results, st.session_state.input_values)
                    word_path = "temp_lightning.docx"
                    word.save(word_path)
                    with open(word_path, "rb") as f:
                        word_bytes = f.read()
                    b64 = base64.b64encode(word_bytes).decode()
                    filename = f"Lightning_Report_{format_pakistan_date()}.docx"
                    os.remove(word_path)
                    st.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="download-btn">📥 Click here to download word report</a>', unsafe_allow_html=True)
                    st.success("✅ Word generated successfully!")
                except Exception as e:
                    st.error(f"Error generating word document: {str(e)}")

# ========== CABLE SIZING TAB ==========
elif st.session_state.selected_calculator == "Cable Sizing":
    st.markdown('<div class="report-header">🔌 Cable sizing calculator</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        efficiency_value = st.number_input("Motor efficiency", value=1.0, min_value=0.5, max_value=1.0, step=0.05, format="%.2f")
        st.caption("1.0 = 100%, 0.95 = 95%")
    with col2:
        ambient_temp = st.number_input("Ambient temperature (°C) - common", value=30.0, min_value=10.0, max_value=80.0, step=5.0, key="ambient_temp_global")
        st.info("Same for all cables")
    
    clearance_options_dict = {
        'touching': 'Touching (0 clearance)',
        'one_diameter': 'One cable diameter spacing',
        'clearance_0_125m': '0.125 m clearance',
        'clearance_0_25m': '0.25 m clearance',
        'clearance_0_5m': '0.5 m clearance',
        'clearance_1_0m': '1.0 m clearance'
    }
    
    if st.button("📥 Import loads from load sheet", use_container_width=True):
        new_loads = []
        for idx, load in st.session_state.universal_loads.iterrows():
            if load['Voltage (V)'] > 300:
                phase = '3-phase'
            else:
                phase = '1-phase'
            load_type_diversity = LOAD_TYPE_FACTORS[load['Load Type']]['diversity']
            
            if load['Voltage (V)'] <= 1000:
                cable_type = 'multi_core_non_armoured'
                install_method = 'C'
                table_config = 'C34' if phase == '3-phase' else 'C2'
                arrangement = 'bunched_in_air_surface_enclosed'
                formation = 'flat'
            else:
                cable_type = 'single_core_non_armoured'
                install_method = 'F'
                table_config = 'F34_flat' if phase == '3-phase' else 'F2'
                arrangement = 'single_layer_ladder_cleats'
                formation = 'flat'
            
            new_loads.append({
                'Load Name': load['Load Description'],
                'Power (kW)': load['Rating (kW)'] * load['Quantity'] * load_type_diversity,
                'Voltage (V)': load['Voltage (V)'],
                'Phase': phase,
                'Load Type': load['Load Type'],
                'Power Factor': load['Power Factor'],
                'Efficiency': efficiency_value,
                'Length (m)': 50,
                'Insulation Type': 'XLPE_90',
                'Cable Type': cable_type,
                'Installation Method': install_method,
                'Table_Config': table_config,
                'Cables in Group': 3,
                'Cable Arrangement': arrangement,
                'Cable Formation': formation,
                'Cable Clearance': 'touching',
                'Soil Resistivity (K.m/W)': 1.5,
                'Burial Depth (m)': 0.8
            })
        st.session_state.loads_df = pd.DataFrame(new_loads)
        st.success("✅ Loads imported successfully!")
        st.rerun()
    
    cable_calc = CableSizingCalculator()
    cable_tabs = st.tabs(["📥 Loads and derating input", "📊 Derating factors summary", "🔌 Cable selection", "⚡ Circuit breakers", "📥 Download report"])
    
    with cable_tabs[0]:
        st.markdown("### 📋 Input parameters")
        if st.session_state.loads_df.empty:
            st.info("No loads imported yet. Click 'Import loads from load sheet' above.")
        else:
            edited_df = st.session_state.loads_df.copy()
            
            st.markdown("#### Edit Load Parameters")
            
            for idx, load in edited_df.iterrows():
                with st.container():
                    st.markdown(f"**Load {idx+1}: {load['Load Name']}**")
                    
                    col_a, col_b, col_c = st.columns(3)
                    
                    with col_a:
                        new_length = st.number_input(
                            f"Length (m)",
                            value=float(load['Length (m)']),
                            min_value=1.0,
                            max_value=5000.0,
                            step=10.0,
                            key=f"length_{idx}"
                        )
                    
                    with col_b:
                        new_insulation = st.selectbox(
                            f"Insulation",
                            options=['XLPE_90', 'PVC_70'],
                            index=0 if load['Insulation Type'] == 'XLPE_90' else 1,
                            format_func=format_insulation_type,
                            key=f"insulation_{idx}"
                        )
                    
                    with col_c:
                        new_cable_type = st.selectbox(
                            f"Cable Type",
                            options=['single_core_non_armoured', 'multi_core_non_armoured', 
                                    'single_core_armoured', 'multi_core_armoured'],
                            index=['single_core_non_armoured', 'multi_core_non_armoured', 
                                   'single_core_armoured', 'multi_core_armoured'].index(load['Cable Type']),
                            format_func=format_cable_type,
                            key=f"cable_type_{idx}"
                        )
                    
                    valid_methods = get_valid_reference_methods(new_cable_type)
                    current_method = load['Installation Method']
                    if current_method not in valid_methods:
                        current_method = valid_methods[0]
                    
                    col_d, col_e, col_f = st.columns(3)
                    
                    with col_d:
                        new_install_method = st.selectbox(
                            f"Reference Method",
                            options=valid_methods,
                            index=valid_methods.index(current_method),
                            format_func=format_installation_method,
                            key=f"install_{idx}"
                        )
                    
                    configs = get_table_configurations(new_cable_type, new_install_method)
                    current_config = load.get('Table_Config', configs[0]['key'] if configs else '')
                    if current_config not in [c['key'] for c in configs]:
                        current_config = configs[0]['key'] if configs else ''
                    
                    with col_e:
                        if configs:
                            new_config = st.selectbox(
                                f"Table Configuration",
                                options=[c['key'] for c in configs],
                                index=[c['key'] for c in configs].index(current_config) if current_config in [c['key'] for c in configs] else 0,
                                format_func=lambda x: next((c['description'] for c in configs if c['key'] == x), x),
                                key=f"config_{idx}"
                            )
                            table_ref = get_table_reference_info(new_cable_type, new_install_method, new_config)
                            st.caption(f"📖 Reference: {table_ref}")
                            
                            selected_config = next((c for c in configs if c['key'] == new_config), None)
                            if selected_config:
                                auto_phase = selected_config['phase']
                            else:
                                auto_phase = load['Phase']
                        else:
                            new_config = ''
                            auto_phase = load['Phase']
                            st.warning("No configurations available")
                    
                    with col_f:
                        new_cables_group = st.number_input(
                            f"Cables in Group",
                            value=int(load['Cables in Group']),
                            min_value=1,
                            max_value=20,
                            step=1,
                            key=f"cables_group_{idx}"
                        )
                    
                    col_g, col_h, col_i = st.columns(3)
                    
                    valid_arrangements = get_valid_arrangements(new_install_method, new_cable_type)
                    current_arrangement = load['Cable Arrangement']
                    if current_arrangement not in valid_arrangements:
                        current_arrangement = valid_arrangements[0] if valid_arrangements else 'bunched_in_air_surface_enclosed'
                    
                    with col_g:
                        new_arrangement = st.selectbox(
                            f"Cable Arrangement",
                            options=valid_arrangements,
                            index=valid_arrangements.index(current_arrangement) if current_arrangement in valid_arrangements else 0,
                            format_func=format_cable_arrangement,
                            key=f"arrangement_{idx}"
                        )
                    
                    with col_h:
                        if 'single_core' in new_cable_type and new_install_method in ['F', 'G']:
                            formation_options = ['flat', 'trefoil', 'spaced']
                        elif 'multi_core' in new_cable_type:
                            formation_options = ['flat']
                        else:
                            formation_options = ['flat', 'trefoil', 'spaced']
                        
                        current_formation = load['Cable Formation']
                        if current_formation not in formation_options:
                            current_formation = formation_options[0]
                        
                        new_formation = st.selectbox(
                            f"Cable Formation",
                            options=formation_options,
                            index=formation_options.index(current_formation),
                            format_func=format_cable_formation,
                            key=f"formation_{idx}"
                        )
                    
                    with col_i:
                        is_buried_valid = (new_install_method in ['D', 'D_direct'] and 'multi_core_armoured' in new_cable_type)
                        
                        if is_buried_valid:
                            clearance_options = get_clearance_options(new_install_method, new_arrangement, 'single_core' in new_cable_type)
                            current_clearance = load.get('Cable Clearance', 'touching')
                            if current_clearance not in clearance_options:
                                current_clearance = clearance_options[0] if clearance_options else 'touching'
                            
                            new_clearance = st.selectbox(
                                f"Cable Clearance",
                                options=clearance_options,
                                index=clearance_options.index(current_clearance) if current_clearance in clearance_options else 0,
                                format_func=lambda x: clearance_options_dict.get(x, x),
                                key=f"clearance_{idx}"
                            )
                            
                            new_soil_res = st.number_input(
                                f"Soil Resistivity (K.m/W)",
                                value=float(load.get('Soil Resistivity (K.m/W)', 1.5)),
                                min_value=0.5,
                                max_value=3.0,
                                step=0.1,
                                key=f"soil_{idx}"
                            )
                            new_depth = st.number_input(
                                f"Burial Depth (m)",
                                value=float(load.get('Burial Depth (m)', 0.8)),
                                min_value=0.3,
                                max_value=3.0,
                                step=0.1,
                                key=f"depth_{idx}"
                            )
                        else:
                            new_clearance = 'touching'
                            new_soil_res = 1.5
                            new_depth = 0.8
                            if new_install_method in ['D', 'D_direct']:
                                st.info(f"ℹ️ Burial parameters only applicable for Multi-core Armoured cables")
                            st.selectbox(
                                f"Cable Clearance (N/A)",
                                options=['touching'],
                                disabled=True,
                                key=f"clearance_disabled_{idx}"
                            )
                            st.number_input(
                                f"Soil Resistivity (N/A)",
                                value=1.5,
                                disabled=True,
                                key=f"soil_disabled_{idx}"
                            )
                            st.number_input(
                                f"Burial Depth (N/A)",
                                value=0.8,
                                disabled=True,
                                key=f"depth_disabled_{idx}"
                            )
                    
                    edited_df.at[idx, 'Length (m)'] = new_length
                    edited_df.at[idx, 'Insulation Type'] = new_insulation
                    edited_df.at[idx, 'Cable Type'] = new_cable_type
                    edited_df.at[idx, 'Installation Method'] = new_install_method
                    edited_df.at[idx, 'Table_Config'] = new_config
                    edited_df.at[idx, 'Phase'] = auto_phase
                    edited_df.at[idx, 'Cables in Group'] = new_cables_group
                    edited_df.at[idx, 'Cable Arrangement'] = new_arrangement
                    edited_df.at[idx, 'Cable Formation'] = new_formation
                    edited_df.at[idx, 'Cable Clearance'] = new_clearance if is_buried_valid else 'touching'
                    edited_df.at[idx, 'Soil Resistivity (K.m/W)'] = new_soil_res if is_buried_valid else 1.5
                    edited_df.at[idx, 'Burial Depth (m)'] = new_depth if is_buried_valid else 0.8
                    
                    st.markdown("---")
            
            st.session_state.loads_df = edited_df
            
            has_error = False
            for idx, load in st.session_state.loads_df.iterrows():
                valid_methods = get_valid_reference_methods(load['Cable Type'])
                if load['Installation Method'] not in valid_methods:
                    st.error(f"❌ Load {idx+1} ({load['Load Name']}): {format_cable_type(load['Cable Type'])} cannot use Method {load['Installation Method']}. Valid methods: {', '.join(valid_methods)}")
                    has_error = True
            
            if not has_error:
                if st.button("🔧 Calculate with derating factors (auto selection)", type="primary", use_container_width=True):
                    with st.spinner("Calculating with automatic cable selection..."):
                        cable_results = []
                        detailed_calcs = []
                        all_factors = {}
                        for idx, load in st.session_state.loads_df.iterrows():
                            cable_category, _ = cable_calc.get_cable_category(load['Voltage (V)'])
                            insulation_type = load['Insulation Type']
                            insulation_temp = 90 if insulation_type == 'XLPE_90' else 70
                            
                            current = cable_calc.calculate_load_current(
                                load['Power (kW)'], load['Voltage (V)'], load['Power Factor'], 
                                load['Efficiency'], load['Phase']
                            )
                            
                            selected_size, cable_data, base_amp, derated_amp, vd_pct, total_k, factors, success, _ = select_cable_automatically(
                                load, cable_calc, ambient_temp,
                                insulation_temp, current,
                                load['Length (m)'], load['Power Factor'], load['Voltage (V)'], load['Phase'],
                                load['Installation Method'], load['Cable Formation'], load['Cable Type'],
                                load['Cable Arrangement'],
                                load['Soil Resistivity (K.m/W)'], load['Burial Depth (m)'], load['Cables in Group'],
                                load['Table_Config'], load.get('Cable Clearance', 'touching')
                            )
                            
                            if selected_size is None:
                                st.error(f"No suitable cable found for {load['Load Name']}")
                                continue
                            
                            all_factors[load['Load Name']] = factors
                            insulation_short = 'PVC' if insulation_type == 'PVC_70' else 'XLPE'
                            base_ampacity = base_amp
                            is_single_core = (load['Cable Type'] in ['single_core_non_armoured', 'single_core_armoured'])
                            total_k_actual, factors_actual = cable_calc.get_derating_factors(
                                ambient_temp, insulation_temp,
                                load['Cables in Group'], load['Cable Arrangement'],
                                load['Installation Method'],
                                load['Soil Resistivity (K.m/W)'], load['Burial Depth (m)'],
                                is_single_core, load.get('Cable Clearance', 'touching')
                            )
                            derated_amp_actual = base_ampacity * total_k_actual
                            isc, k_value, operating_temp, theta_i, theta_f = cable_calc.calculate_short_circuit(
                                selected_size, insulation_short,
                                ambient_temp, current, base_ampacity,
                                factors_actual['k1 (Temperature)'],
                                factors_actual['k2 (Grouping)'],
                                factors_actual['k3 (Soil Resistivity)'],
                                factors_actual['k4 (Depth)'],
                                1.0
                            )
                            status = 'PASS' if success else 'FAIL'
                            cable_results.append({
                                'Load Name': load['Load Name'],
                                'Load Type': load.get('Load Type', 'Continuous'),
                                'Power (kW)': load['Power (kW)'],
                                'Voltage (V)': load['Voltage (V)'],
                                'Phase': load['Phase'],
                                'PF': load['Power Factor'],
                                'Efficiency': f"{load['Efficiency']*100:.0f}%",
                                'Length (m)': load['Length (m)'],
                                'Insulation': format_insulation_type(insulation_type),
                                'Cable Category': cable_category,
                                'Cable Type': format_cable_type(load['Cable Type']),
                                'Size (mm²)': selected_size,
                                'Load Current (A)': f"{current:.1f}",
                                'Current Carrying Capacity (A)': base_ampacity,
                                'Derating Factor K': f"{total_k_actual:.3f}",
                                'Derated Ampacity (A)': f"{derated_amp_actual:.1f}",
                                'Voltage Drop (%)': f"{vd_pct:.3f}",
                                'Short Circuit (kA)': f"{isc/1000:.2f}",
                                'Status': status,
                                'Check': 'PASS' if (vd_pct <= 2.5 and derated_amp_actual >= current) else 'FAIL'
                            })
                            detailed_calcs.append({
                                'load_name': load['Load Name'],
                                'load_type': load.get('Load Type', 'Continuous'),
                                'power': load['Power (kW)'],
                                'voltage': load['Voltage (V)'],
                                'phase': load['Phase'],
                                'pf': load['Power Factor'],
                                'efficiency': load['Efficiency'],
                                'length': load['Length (m)'],
                                'current': current,
                                'size': selected_size,
                                'insulation_type': insulation_type,
                                'cable_category': cable_category,
                                'cable_type': load['Cable Type'],
                                'formation': load['Cable Formation'],
                                'installation': load['Installation Method'],
                                'arrangement': load['Cable Arrangement'],
                                'clearance': load.get('Cable Clearance', 'touching'),
                                'soil_res': load['Soil Resistivity (K.m/W)'],
                                'depth': load['Burial Depth (m)'],
                                'num_cables': load['Cables in Group'],
                                'base_amp': base_ampacity,
                                'derated_amp': derated_amp_actual,
                                'vd_pct': vd_pct,
                                'sc': isc/1000,
                                'theta_i': theta_i,
                                'theta_f': theta_f,
                                'operating_temp': operating_temp,
                                'k1': factors_actual['k1 (Temperature)'],
                                'k2': factors_actual['k2 (Grouping)'],
                                'k3': factors_actual['k3 (Soil Resistivity)'],
                                'k4': factors_actual['k4 (Depth)'],
                                'total_k': total_k_actual,
                                'ambient_temp': ambient_temp,
                                'status': status,
                                'vd_pass': vd_pct <= 2.5,
                                'ampacity_pass': derated_amp_actual >= current,
                                'trials': []
                            })
                            if not success:
                                st.warning(f"⚠️ {load['Load Name']}: Even largest cable {selected_size} mm² fails! vd={vd_pct:.2f}% > 2.5%")
                            else:
                                st.success(f"✅ {load['Load Name']}: Selected {selected_size} mm² cable (vd={vd_pct:.2f}%)")
                        
                        st.session_state.cable_results_df = pd.DataFrame(cable_results)
                        st.session_state.detailed_calcs = detailed_calcs
                        st.session_state.all_derating_factors = all_factors
                        cb_calc = CircuitBreakerCalculator()
                        manufacturer = 'Schneider Electric'
                        cb_results, cb_details = cb_calc.calculate_cb_size(
                            st.session_state.loads_df, 1.25, manufacturer
                        )
                        main_cbs_by_voltage, main_cb_details_by_voltage = cb_calc.calculate_main_cb_by_voltage(
                            st.session_state.loads_df, 1.25
                        )
                        st.session_state.cb_results = cb_results
                        st.session_state.cb_details = cb_details
                        st.session_state.main_cbs_by_voltage = main_cbs_by_voltage
                        st.session_state.main_cb_details_by_voltage = main_cb_details_by_voltage
                        st.success("✅ Calculations complete with automatic cable selection!")
            else:
                st.warning("⚠️ Please fix the errors above before calculating.")
    
    with cable_tabs[1]:
        st.markdown('<div class="report-header">Derating factors summary</div>', unsafe_allow_html=True)
        st.markdown("""
### Derating Factor Formulas
**Total derating factor K = k1 × k2 × k3 × k4**

**k1 (Temperature correction)** - Based on IEC 60364-5-52 Table 4B1/4B2
**k2 (Grouping correction)** - Based on IEC 60364-5-52 Table 4C1 (for air) or Table 4C2/4C3 (for buried)
**k3 (Soil resistivity correction)** - Based on IEC 60502-2 Tables B.14, B.15, B.16 (Reference: 1.5 K.m/W)
**k4 (Depth correction)** - Based on IEC 60502-2 Tables B.12, B.13 (Reference: 0.8m)

**Derated ampacity = Current Carrying Capacity × K**
""")
        if hasattr(st.session_state, 'all_derating_factors') and st.session_state.all_derating_factors:
            for load_name, factors in st.session_state.all_derating_factors.items():
                with st.expander(f"📊 {load_name} - Derating factors"):
                    st.markdown(f"""
| Factor | Value | Description |
|--------|-------|-------------|
| k1 (Temperature) | {factors['k1 (Temperature)']:.3f} | Based on ambient temperature and insulation type |
| k2 (Grouping) | {factors['k2 (Grouping)']:.3f} | Based on number of circuits and installation arrangement |
| k3 (Soil resistivity) | {factors['k3 (Soil Resistivity)']:.3f} | Based on IEC 60502-2 Tables B.14-B.16 |
| k4 (Depth) | {factors['k4 (Depth)']:.3f} | Based on IEC 60502-2 Tables B.12-B.13 |
| **Total K** | **{factors['total']:.3f}** | **K = k1 × k2 × k3 × k4** |
""")
        else:
            st.info("👈 Calculate loads first to see derating factors")
    
    with cable_tabs[2]:
        st.markdown('<div class="report-header">🔌 Cable selection results</div>', unsafe_allow_html=True)
        st.markdown("### ⚡ Voltage drop limit: **2.5%**")
        if not st.session_state.cable_results_df.empty:
            st.dataframe(st.session_state.cable_results_df, use_container_width=True, hide_index=True)
            st.markdown("### 📋 Detailed calculation")
            for calc in st.session_state.detailed_calcs:
                with st.expander(f"🔍 {calc['load_name']} ({format_load_type(calc.get('load_type', 'Continuous'))})"):
                    st.markdown(f"""
**Step 1: Load current**  
I = {calc['power']:.1f} x 1000 / (1.732 x {calc['voltage']:.0f} x {calc['pf']:.2f} x {calc.get('efficiency', 1.0):.2f}) = **{calc['current']:.1f} A**

**Step 2: Cable type and reference selection**  
Voltage {calc['voltage']:.0f}V -> {calc['cable_category']}  
Cable Type: {format_cable_type(calc['cable_type'])}  
Formation: {format_cable_formation(calc['formation'])}  
Reference Method: {format_installation_method(calc['installation'])}

**Step 3: Derating factors (FOR AMPACITY ONLY)**
- Ambient temperature: {calc['ambient_temp']}°C
- Insulation type: {format_insulation_type(calc['insulation_type'])}
- Cable type: {format_cable_type(calc['cable_type'])}
- Installation method: {calc['installation']}
- Cables in group: {calc['num_cables']}
- Cable arrangement: {format_cable_arrangement(calc['arrangement'])}
- Cable clearance: {format_cable_clearance(calc.get('clearance', 'touching'))}
- Soil resistivity: {calc['soil_res']} K.m/W
- Burial depth: {calc['depth']} m

**Calculated factors:**
- k1 (Temperature): {calc['k1']:.3f}
- k2 (Grouping): {calc['k2']:.3f}
- k3 (Soil resistivity): {calc['k3']:.3f}
- k4 (Depth): {calc['k4']:.3f}
- **Total K = {calc['total_k']:.3f}**

**Step 4: Cable selection**  
Selected: {calc['size']} mm² {format_cable_type(calc['cable_type'])}  
Current Carrying Capacity: {calc['base_amp']} A  
Derated ampacity = CCC × K = {calc['base_amp']} × {calc['total_k']:.3f} = **{calc['derated_amp']:.1f} A**  
Ampacity Check: {calc['derated_amp']:.1f} A >= {calc['current']:.1f} A -> **{'PASS' if calc['ampacity_pass'] else 'FAIL'}**

**Step 5: Voltage drop calculation (NO DERATING FACTORS)**
""")
                    
                    vd_values = get_voltage_drop_values(calc['cable_type'], calc['size'], calc['phase'], calc['formation'])
                    
                    if vd_values['type'] == 'mv':
                        mv_value = vd_values['value']
                        st.markdown(f"""
Using mV/A/m = {mv_value} from BS 7671 Table

Formula: Vd = mV/A/m × I × L / 1000
Vd = {mv_value} × {calc['current']:.1f} × {calc['length']:.0f} / 1000
Vd = **{mv_value * calc['current'] * calc['length'] / 1000:.2f} V**
Vd% = ({mv_value * calc['current'] * calc['length'] / 1000:.2f} / {calc['voltage']:.0f}) × 100 = **{calc['vd_pct']:.3f}%**
""")
                    else:
                        r = vd_values['R']
                        x = vd_values['X']
                        phi = math.acos(calc['pf'])
                        sin_phi = math.sin(phi)
                        
                        st.markdown(f"""
Using R = {r:.4f} Ω/km, X = {x:.4f} Ω/km from BS 7671 Table
Power factor cosφ = {calc['pf']:.3f}, sinφ = √(1 - cos²φ) = {sin_phi:.4f}
R cosφ = {r:.4f} × {calc['pf']:.3f} = {r * calc['pf']:.4f}
X sinφ = {x:.4f} × {sin_phi:.4f} = {x * sin_phi:.4f}
(R cosφ + X sinφ) = {r * calc['pf'] + x * sin_phi:.4f}
""")
                        if calc['phase'] == '3-phase':
                            st.markdown(f"""
Formula (3-phase): Vd = √3 × I × (R cosφ + X sinφ) × L / 1000
Vd = 1.732 × {calc['current']:.1f} × {(r * calc['pf'] + x * sin_phi):.4f} × {calc['length']:.0f} / 1000
Vd = **{calc['vd_pct'] * calc['voltage'] / 100:.2f} V**
Vd% = **{calc['vd_pct']:.3f}%**
""")
                        else:
                            st.markdown(f"""
Formula (1-phase): Vd = 2 × I × (R cosφ + X sinφ) × L / 1000
Vd = 2 × {calc['current']:.1f} × {(r * calc['pf'] + x * sin_phi):.4f} × {calc['length']:.0f} / 1000
Vd = **{calc['vd_pct'] * calc['voltage'] / 100:.2f} V**
Vd% = **{calc['vd_pct']:.3f}%**
""")
                    
                    st.markdown(f"""
**Result:**  
Voltage drop = **{calc['vd_pct']:.3f}%** (Limit: 2.5%)  
VD Check: {calc['vd_pct']:.3f}% <= 2.5% -> **{'PASS' if calc['vd_pass'] else 'FAIL'}**

**Step 6: Short circuit calculation**  
Isc = **{calc['sc']:.2f} kA**

**Final status: {'PASS' if calc['status'] == 'PASS' else 'FAIL'}**
""")
        else:
            st.info("👈 Calculate loads first")
    
    with cable_tabs[3]:
        st.markdown('<div class="report-header">⚡ Circuit breaker sizing</div>', unsafe_allow_html=True)
        
        st.markdown("### 🏭 Manufacturer Selection")
        manufacturer_options = list(MANUFACTURERS.keys())
        selected_manufacturer = st.selectbox("Select Circuit Breaker Manufacturer", options=manufacturer_options, index=0, key="cb_manufacturer_select")
        st.info(f"**Selected Manufacturer:** {selected_manufacturer}\n\n**Series:** MCB: {MANUFACTURERS[selected_manufacturer]['MCB']} | MCCB: {MANUFACTURERS[selected_manufacturer]['MCCB']} | ACB: {MANUFACTURERS[selected_manufacturer]['ACB']}")
        
        if st.session_state.cb_results:
            for i in range(len(st.session_state.cb_results)):
                st.session_state.cb_results[i]['Manufacturer'] = selected_manufacturer
                if st.session_state.cb_results[i]['Breaker Type'] in MANUFACTURERS[selected_manufacturer]:
                    st.session_state.cb_results[i]['Series'] = MANUFACTURERS[selected_manufacturer][st.session_state.cb_results[i]['Breaker Type']]
                else:
                    st.session_state.cb_results[i]['Series'] = 'Standard series'
            for i in range(len(st.session_state.cb_details)):
                st.session_state.cb_details[i]['manufacturer'] = selected_manufacturer
                if st.session_state.cb_details[i]['breaker_type'] in MANUFACTURERS[selected_manufacturer]:
                    st.session_state.cb_details[i]['series'] = MANUFACTURERS[selected_manufacturer][st.session_state.cb_details[i]['breaker_type']]
                else:
                    st.session_state.cb_details[i]['series'] = 'Standard series'
            
            st.markdown("### ⚡ Individual circuit breakers")
            st.markdown("#### 🔧 Pole selection for each load")
            st.info("Select the number of poles based on your system requirements:")
            pole_options = {
                '1P': 'Single pole - Phase only (for IT systems or DC)',
                '2P': 'Two pole - Phase + neutral (for single-phase TN/TT systems)',
                '3P': 'Three pole - 3 phases only (for TN-C or 3-wire systems)',
                '4P': 'Four pole - 3 phases + neutral (for TN-S systems)'
            }
            pole_selections = {}
            cols = st.columns(min(len(st.session_state.cb_results), 3))
            for idx, result in enumerate(st.session_state.cb_results):
                col_idx = idx % 3
                with cols[col_idx]:
                    default_poles = '3P' if result['Phase'] == '3-phase' else ('2P' if result['Phase'] == '1-phase' else '1P')
                    unique_key = f"pole_{result['Load']}_{idx}"
                    pole_selections[result['Load']] = st.selectbox(
                        f"Poles for {result['Load']}",
                        options=list(pole_options.keys()),
                        format_func=lambda x: f"{x} - {pole_options[x].split('-')[0].strip()}",
                        key=unique_key,
                        index=list(pole_options.keys()).index(default_poles)
                    )
                    st.caption(pole_options[pole_selections[result['Load']]])
            
            cb_df = pd.DataFrame([{
                'Load': r['Load'],
                'Power (kW)': r['Power (kW)'],
                'Voltage (V)': r['Voltage (V)'],
                'Current (A)': f"{r['Current (A)']:.1f}",
                'Required (A)': f"{r['Required CB (A)']:.1f}",
                'Selected CB (A)': r['Selected CB (A)'],
                'Type': f"{r['Breaker Type']}",
                'Manufacturer': selected_manufacturer,
                'Series': r['Series'],
                'Poles': pole_selections.get(r['Load'], '3P'),
                'Standard': r['Standard']
            } for r in st.session_state.cb_results])
            st.dataframe(cb_df, use_container_width=True, hide_index=True)
            
            st.markdown("### 🔋 Main Circuit Breakers (by Voltage Level)")
            st.info("Main circuit breakers are calculated separately for each voltage level based on the loads connected to that system.")
            
            main_pole_selections = {}
            if st.session_state.main_cbs_by_voltage:
                for voltage, main_cb in st.session_state.main_cbs_by_voltage.items():
                    with st.expander(f"📌 {main_cb['system_type']} - {main_cb['voltage_range']} System", expanded=True):
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Total Power", f"{main_cb['total_power']:.1f} kW")
                            st.metric("Number of Loads", main_cb['num_loads'])
                        with col2:
                            st.metric("Total Current", f"{main_cb['current']:.1f} A")
                            st.metric("Average PF", f"{main_cb['avg_pf']:.3f}")
                        with col3:
                            st.metric("Required CB", f"{main_cb['required_cb']:.1f} A")
                            st.metric("Selected CB", f"{main_cb['selected_cb']} A")
                        st.markdown(f"**Breaker Type:** {main_cb['breaker_type']} ({main_cb['standard']})")
                        
                        default_main_poles = '3P'
                        main_pole_selections[f'voltage_{voltage}'] = st.selectbox(
                            f"Poles for {main_cb['voltage_range']} Main Breaker",
                            options=list(pole_options.keys()),
                            format_func=lambda x: f"{x} - {pole_options[x]}",
                            key=f"main_pole_voltage_{voltage}",
                            index=list(pole_options.keys()).index(default_main_poles)
                        )
                
                st.markdown("#### 📊 Main Circuit Breakers Summary")
                main_cb_summary = []
                for voltage, main_cb in st.session_state.main_cbs_by_voltage.items():
                    main_cb_summary.append({
                        'System': f"{main_cb['voltage_range']} ({main_cb['system_type']})",
                        'Total Power (kW)': f"{main_cb['total_power']:.1f}",
                        'Total Current (A)': f"{main_cb['current']:.1f}",
                        'Selected CB (A)': main_cb['selected_cb'],
                        'Breaker Type': main_cb['breaker_type'],
                        'Poles': main_pole_selections.get(f'voltage_{voltage}', '3P')
                    })
                if main_cb_summary:
                    st.dataframe(pd.DataFrame(main_cb_summary), use_container_width=True, hide_index=True)
            else:
                st.info("No main circuit breaker calculations available.")
        else:
            st.info("👈 Calculate cable sizes first to see circuit breaker results")
    
    with cable_tabs[4]:
        st.markdown('<div class="report-header">📥 Download report</div>', unsafe_allow_html=True)
        if not st.session_state.cable_results_df.empty and st.session_state.cb_results:
            if st.button("📥 Generate word report", key="cable_word", use_container_width=True):
                with st.spinner("Generating word with complete detailed calculations..."):
                    try:
                        for calc in st.session_state.detailed_calcs:
                            if 'pf' not in calc or calc['pf'] is None:
                                calc['pf'] = 0.85
                            if 'formation' not in calc:
                                calc['formation'] = 'flat'
                        
                        word = CableWordReport()
                        word.add_title()
                        word.add_common_parameters(ambient_temp)
                        word.add_load_details(st.session_state.loads_df)
                        word.add_cable_results(st.session_state.cable_results_df)
                        
                        if st.session_state.detailed_calcs:
                            word.add_detailed_calculations(st.session_state.detailed_calcs)
                        
                        if st.session_state.cb_results and st.session_state.main_cbs_by_voltage:
                            pole_selections = {}
                            for r in st.session_state.cb_results:
                                if r['Phase'] == '3-phase':
                                    pole_selections[r['Load']] = '3P'
                                elif r['Phase'] == '1-phase':
                                    pole_selections[r['Load']] = '2P'
                                else:
                                    pole_selections[r['Load']] = '1P'
                            
                            main_pole_selections = {}
                            for voltage in st.session_state.main_cbs_by_voltage.keys():
                                main_pole_selections[f'voltage_{voltage}'] = '3P'
                            
                            word.add_cb_results(
                                st.session_state.cb_results, 
                                st.session_state.main_cbs_by_voltage, 
                                pole_selections, 
                                main_pole_selections,
                                st.session_state.cb_details if st.session_state.cb_details else [],
                                selected_manufacturer
                            )
                        
                        word_path = "temp_cable_report.docx"
                        word.save(word_path)
                        
                        with open(word_path, "rb") as f:
                            word_bytes = f.read()
                        
                        b64 = base64.b64encode(word_bytes).decode()
                        filename = f"Cable_CB_Report_{format_pakistan_date()}.docx"
                        os.remove(word_path)
                        
                        st.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="download-btn">📥 Click here to download word report</a>', unsafe_allow_html=True)
                        st.success("✅ Word generated successfully!")
                        
                    except Exception as e:
                        st.error(f"Error generating word document: {str(e)}")
                        st.code(traceback.format_exc())
        else:
            st.info("👈 Calculate cable sizes first to generate report")

# ========== TRANSFORMER SIZING TAB ==========
elif st.session_state.selected_calculator == "Transformer Sizing":
    st.markdown('<div class="report-header">⚙️ TRANSFORMER SIZING CALCULATOR</div>', unsafe_allow_html=True)
    st.markdown(f"""
    <div class="info-box">
        <h4>📌 Using loads from universal load sheet</h4>
        <p>Total {len(st.session_state.universal_loads)} loads available. Calculations below use these loads.</p>
    </div>
    """, unsafe_allow_html=True)
    
    tx_main_tabs = st.tabs(["📊 Load analysis", "📈 Largest equipment analysis", "📥 Download report"])
    tx_calc = SimpleTransformerCalculator()
    
    with tx_main_tabs[0]:
        load_sub_tabs = st.tabs(["📋 Step-by-step p, q, s", "📊 Summary table"])
        with load_sub_tabs[0]:
            st.markdown("### 📋 Step-by-step calculations for each load")
            total_p = 0
            total_q = 0
            for idx, load in st.session_state.universal_loads.iterrows():
                connected = load['Rating (kW)'] * load['Quantity']
                load_type_diversity = LOAD_TYPE_FACTORS[load['Load Type']]['diversity']
                p = tx_calc.calculate_p(load['Rating (kW)'], load['Quantity'], load_type_diversity)
                phi = math.acos(load['Power Factor'])
                tan_phi = math.tan(phi)
                q = tx_calc.calculate_q(p, load['Power Factor'])
                s = tx_calc.calculate_s(p, q)
                cable_type = "LV" if load['Voltage (V)'] <= 1000 else "MV"
                st.markdown(f"""
                <div class="calc-step">
                    <h4>📌 Load {idx+1}: <span style="color: #1E3A8A;">{load['Load Description']}</span> ({load['Load Type']}, {cable_type})</h4>
                    <p><b>Input parameters:</b> Rating = {load['Rating (kW)']:.0f} kW, Quantity = {load['Quantity']}, Pf = {load['Power Factor']}, Load type = {load['Load Type']} ({load_type_diversity*100:.0f}%)</p>
                    <p><b>Step 1 - Connected power:</b> {load['Rating (kW)']:.0f} kW x {load['Quantity']} = <b>{connected:.0f} kW</b></p>
                    <p><b>Step 2 - Demand power (P):</b> {connected:.0f} kW x {load_type_diversity:.1f} = <b>{p:.1f} kW</b></p>
                    <p><b>Step 3 - Angle φ:</b> acos({load['Power Factor']}) = <b>{math.degrees(phi):.1f}°</b></p>
                    <p><b>Step 4 - tan(φ):</b> tan({math.degrees(phi):.1f}°) = <b>{tan_phi:.3f}</b></p>
                    <p><b>Step 5 - Reactive power (Q):</b> {p:.1f} kW x {tan_phi:.3f} = <b>{q:.1f} kVAR</b></p>
                    <p><b>Step 6 - Apparent power (S):</b> √({p:.1f}² + {q:.1f}²) = <b>{s:.1f} kVA</b></p>
                </div>
                """, unsafe_allow_html=True)
                total_p += p
                total_q += q
            st.session_state.total_p = total_p
            st.session_state.total_q = total_q
        
        with load_sub_tabs[1]:
            st.markdown("### 📊 Load summary table")
            summary_data = []
            for idx, load in st.session_state.universal_loads.iterrows():
                load_type_diversity = LOAD_TYPE_FACTORS[load['Load Type']]['diversity']
                connected = load['Rating (kW)'] * load['Quantity']
                p = tx_calc.calculate_p(load['Rating (kW)'], load['Quantity'], load_type_diversity)
                q = tx_calc.calculate_q(p, load['Power Factor'])
                s = tx_calc.calculate_s(p, q)
                cable_type = "LV" if load['Voltage (V)'] <= 1000 else "MV"
                summary_data.append({
                    'Load': load['Load Description'],
                    'Type': load['Load Type'],
                    'Cable': cable_type,
                    'Qty': load['Quantity'],
                    'Rating (kW)': load['Rating (kW)'],
                    'Connected (kW)': f"{connected:.0f}",
                    'Load factor': f"{load_type_diversity*100:.0f}%",
                    'P (kW)': f"{p:.1f}",
                    'Pf': load['Power Factor'],
                    'Q (kVAR)': f"{q:.1f}",
                    'S (kVA)': f"{s:.1f}"
                })
            summary_df = pd.DataFrame(summary_data)
            st.dataframe(summary_df, use_container_width=True, hide_index=True)
            st.markdown("""
            <div class="formula-box">
                <h4>📐 Formulas used:</h4>
                <p><b>Connected power = Rating x Quantity</b></p>
                <p><b>P = Connected power x Load type factor</b> (Real power)</p>
                <p><b>Q = P x tan(acos(PF))</b> (Reactive power)</p>
                <p><b>S = √(P² + Q²)</b> (Apparent power)</p>
                <p><b>Load type factors:</b> Continuous=100%, Intermittent=30%, Standby=10%</p>
            </div>
            """, unsafe_allow_html=True)
    
    with tx_main_tabs[1]:
        st.markdown("### 🏭 Largest equipment analysis")
        largest_idx, largest_load, largest_connected = tx_calc.find_largest_equipment(st.session_state.universal_loads)
        if largest_load is not None:
            load_type_diversity = LOAD_TYPE_FACTORS[largest_load['Load Type']]['diversity']
            p_largest = tx_calc.calculate_p(largest_load['Rating (kW)'], largest_load['Quantity'], load_type_diversity)
            q_largest = tx_calc.calculate_q(p_largest, largest_load['Power Factor'])
            s_largest = tx_calc.calculate_s(p_largest, q_largest)
            total_p = 0
            total_q = 0
            all_loads_data = []
            for idx, load in st.session_state.universal_loads.iterrows():
                load_div = LOAD_TYPE_FACTORS[load['Load Type']]['diversity']
                p = tx_calc.calculate_p(load['Rating (kW)'], load['Quantity'], load_div)
                q = tx_calc.calculate_q(p, load['Power Factor'])
                s = tx_calc.calculate_s(p, q)
                total_p += p
                total_q += q
                cable_type = "LV" if load['Voltage (V)'] <= 1000 else "MV"
                all_loads_data.append({
                    'load': load['Load Description'],
                    'type': load['Load Type'],
                    'cable': cable_type,
                    'p': p,
                    'q': q,
                    's': s,
                    'is_largest': (idx == largest_idx)
                })
            total_s = math.sqrt(total_p**2 + total_q**2) if (total_p**2 + total_q**2) > 0 else 0
            st.markdown(f"""
            <div class="largest-equipment">
                <h3>🏆 Largest equipment: {largest_load['Load Description']}</h3>
                <table style="width:100%; border-collapse: collapse;">
                    <tr><td style="padding: 10px; font-weight: bold;">Load type: <td style="padding: 10px;"><span class="value">{largest_load['Load Type']} ({load_type_diversity*100:.0f}%)</span></td>
                    <tr><td style="padding: 10px; font-weight: bold;">Connected power: <td style="padding: 10px;"><span class="value">{largest_connected:.0f} kW</span> ({largest_load['Rating (kW)']:.0f} kW x {largest_load['Quantity']})</span></td>
                    <tr><td style="padding: 10px; font-weight: bold;">Demand power (P): <td style="padding: 10px;"><span class="value">{p_largest:.1f} kW</span> (after {load_type_diversity*100:.0f}% factor)</span></td>
                    <tr><td style="padding: 10px; font-weight: bold;">Reactive power (Q): <td style="padding: 10px;"><span class="value">{q_largest:.1f} kVAR</span> (Pf = {largest_load['Power Factor']})</span></table>
                    <tr><td style="padding: 10px; font-weight: bold;">Apparent power (S): <td style="padding: 10px;"><span class="value">{s_largest:.1f} kVA</span></span></tr>
                </table>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("### 📊 Impact on total system")
            impact_data = []
            for data in all_loads_data:
                p_pct = (data['p'] / total_p) * 100 if total_p > 0 else 0
                s_pct = (data['s'] / total_s) * 100 if total_s > 0 else 0
                impact_data.append({
                    'Load': data['load'],
                    'Type': data['type'],
                    'Cable': data['cable'],
                    'P (kW)': f"{data['p']:.1f}",
                    '% of total P': f"{p_pct:.1f}%",
                    'S (kVA)': f"{data['s']:.1f}",
                    '% of total S': f"{s_pct:.1f}%",
                    'Highlight': '🏆 Largest' if data['is_largest'] else ''
                })
            impact_df = pd.DataFrame(impact_data)
            st.dataframe(impact_df, use_container_width=True, hide_index=True)
            st.markdown("### 📈 Cumulative effect analysis")
            sorted_loads = sorted(all_loads_data, key=lambda x: x['p'], reverse=True)
            cumulative_p = 0
            cumulative_data = []
            for i, data in enumerate(sorted_loads):
                cumulative_p += data['p']
                p_pct = (cumulative_p / total_p) * 100 if total_p > 0 else 0
                cumulative_data.append({
                    'Rank': i+1,
                    'Load': data['load'],
                    'Type': data['type'],
                    'Individual P (kW)': f"{data['p']:.1f}",
                    'Cumulative P (kW)': f"{cumulative_p:.1f}",
                    '% of total': f"{p_pct:.1f}%"
                })
            cumulative_df = pd.DataFrame(cumulative_data)
            st.dataframe(cumulative_df, use_container_width=True, hide_index=True)
            st.session_state.tx_largest_data = {
                'load': largest_load,
                'connected': largest_connected,
                'p': p_largest,
                'q': q_largest,
                's': s_largest
            }
    
    with tx_main_tabs[2]:
        st.markdown("### ⚙️ Future expansion")
        future_expansion = st.number_input("Future expansion (%)", value=20, min_value=0, max_value=100, step=5)
        if 'total_p' in st.session_state and 'total_q' in st.session_state:
            total_p = st.session_state.total_p
            total_q = st.session_state.total_q
            total_s = math.sqrt(total_p**2 + total_q**2)
            with_future = total_s * (1 + future_expansion/100)
            selected_kva = tx_calc.get_standard_rating(with_future)
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total P", f"{total_p:.0f} kW")
            with col2:
                st.metric("Total Q", f"{total_q:.0f} kVAR")
            with col3:
                st.metric("Total S", f"{total_s:.0f} kVA")
            with col4:
                st.metric("With future", f"{with_future:.0f} kVA")
            st.markdown(f"""
            <div class="result-card">
                <h3>✅ Final transformer selection</h3>
                <p><b>S = √(P² + Q²) = √({total_p:.0f}² + {total_q:.0f}²) = {total_s:.0f} kVA</b></p>
                <p><b>With {future_expansion}% future = {total_s:.0f} x {1 + future_expansion/100:.2f} = {with_future:.0f} kVA</b></p>
                <p style="font-size: 24px;"><b>Selected: {selected_kva} kVA</b></p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("📥 Download word report", key="tx_word", use_container_width=True):
                with st.spinner("Generating word report..."):
                    try:
                        word = TransformerWordReport()
                        word.add_title()
                        total_p_load = word.add_load_analysis(st.session_state.universal_loads)
                        total_p_step, total_q_step = word.add_step_by_step(st.session_state.universal_loads)
                        if 'tx_largest_data' in st.session_state and st.session_state.tx_largest_data:
                            word.add_largest_equipment(st.session_state.universal_loads, total_p, total_s, st.session_state.tx_largest_data)
                        word.add_transformer_selection(total_p_step, total_q_step, future_expansion, selected_kva)
                        word_path = "temp_transformer_report.docx"
                        word.save(word_path)
                        with open(word_path, "rb") as f:
                            word_bytes = f.read()
                        b64 = base64.b64encode(word_bytes).decode()
                        filename = f"Transformer_Report_{format_pakistan_date()}.docx"
                        os.remove(word_path)
                        st.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="download-btn">📥 Click here to download word report</a>', unsafe_allow_html=True)
                        st.success("✅ Word generated successfully!")
                    except Exception as e:
                        st.error(f"Error generating word document: {e}")
        else:
            st.warning("⚠️ Please go to load analysis tab first to calculate totals.")

# ========== GENERATOR SIZING TAB ==========
elif st.session_state.selected_calculator == "Generator Sizing":
    st.markdown('<div class="report-header">🔄 GENERATOR SIZING CALCULATOR</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class="info-box">
        <h4> Iterative Development</h4>
        <p>Features:</p>
        <ul>
            <li>Load analysis and kVA calculation</li>
            <li>Generator set selection from standard ratings</li>
            <li>Voltage drop calculation for generator cables</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

# ========== EARTHING TAB ==========
elif st.session_state.selected_calculator == "Earthing":
    st.markdown('<div class="report-header">⏚ EARTHING CALCULATOR</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class="info-box">
        <h4> Iterative Development</h4>
        <p>Features:</p>
        <ul>
            <li>Soil resistivity analysis</li>
            <li>Rod, plate, and strip earthing calculations</li>
            <li>Step and touch potential analysis (IEEE Std 80)</li>
            <li>IEC 62305 lightning protection earthing compliance</li>
            <li>Multiple electrode configuration (parallel rods)</li>
            <li>Earth resistance calculation for various electrode types</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")
st.markdown(f"<div style='text-align: center; color: gray; font-size: 16px;'>🔌 CES-Electrical | Version 3.0 | {format_pakistan_datetime()} (Pakistan Time)</div>", unsafe_allow_html=True)